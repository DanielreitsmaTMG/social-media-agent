"""
Tijdelijke variant van create_sharepoint_doc.py.
Slaat Word-documenten lokaal op in output/JAAR-WW/ in plaats van SharePoint.

Geen credentials nodig — werkt direct.

Usage:
    python systems/save_local_docs.py
    python systems/save_local_docs.py --posts-file intermediates/posts_2026-06-05.json
    python systems/save_local_docs.py --client-id DR-001
"""

import argparse
import io
import json
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MONTHS_NL = [
    "januari", "februari", "maart", "april", "mei", "juni",
    "juli", "augustus", "september", "oktober", "november", "december",
]

PLATFORM_COLORS = {
    "instagram": RGBColor(0xE1, 0x30, 0x6C),
    "linkedin":  RGBColor(0x00, 0x77, 0xB5),
    "facebook":  RGBColor(0x18, 0x77, 0xF2),
}

PLATFORM_LABELS = {
    "instagram": "Instagram",
    "linkedin":  "LinkedIn",
    "facebook":  "Facebook",
}


def _format_date_nl(d: date) -> str:
    return f"{d.day} {MONTHS_NL[d.month - 1]}"


def _week_number(d: date) -> int:
    return d.isocalendar()[1]


def _build_docx(bedrijfsnaam: str, posts: dict, week_start: date, week_end: date) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin   = Pt(70)
        section.right_margin  = Pt(70)

    week_nr    = _week_number(week_start)
    title_text = (
        f"{bedrijfsnaam} — Social Media"
        f" | Week {week_nr}"
        f" | {_format_date_nl(week_start)} – {_format_date_nl(week_end)} {week_end.year}"
    )

    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph(f"Gegenereerd op {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    for platform_key in ("instagram", "linkedin", "facebook"):
        platform_posts = posts.get(platform_key, [])
        if not platform_posts:
            continue

        label = PLATFORM_LABELS[platform_key]
        color = PLATFORM_COLORS[platform_key]
        count = len(platform_posts)

        heading = doc.add_heading(
            f"{label} — {count} post{'s' if count > 1 else ''}",
            level=2,
        )
        for run in heading.runs:
            run.font.color.rgb = color

        for post in platform_posts:
            dag_p = doc.add_paragraph()
            dag_run = dag_p.add_run(f"📅 {post.get('dag', '')}")
            dag_run.bold = True
            dag_run.font.size = Pt(11)

            doc.add_paragraph(post.get("caption", ""))

            ht_p = doc.add_paragraph()
            ht_run = ht_p.add_run(post.get("hashtags", ""))
            ht_run.font.color.rgb = color
            ht_run.font.size = Pt(10)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _latest_file(pattern: str) -> Optional[str]:
    files = sorted(Path("intermediates").glob(pattern), reverse=True)
    return str(files[0]) if files else None


def main():
    parser = argparse.ArgumentParser(description="Sla weekposts lokaal op als Word-bestanden")
    parser.add_argument("--posts-file", help="Pad naar posts JSON")
    parser.add_argument("--client-id", help="Verwerk alleen deze klant")
    parser.add_argument("--output-dir", help="Outputmap (default: output/JAAR-WW/)")
    args = parser.parse_args()

    posts_file = args.posts_file or _latest_file("posts_*.json")
    if not posts_file or not Path(posts_file).exists():
        print("Geen posts-bestand gevonden. Voer eerst generate_weekly_posts.py uit.", file=sys.stderr)
        sys.exit(1)

    with open(posts_file, encoding="utf-8") as f:
        all_posts = json.load(f)

    if args.client_id:
        if args.client_id not in all_posts:
            print(f"Klant '{args.client_id}' niet gevonden.", file=sys.stderr)
            sys.exit(1)
        all_posts = {args.client_id: all_posts[args.client_id]}

    # Bepaal outputmap op basis van weeknummer
    first = next(iter(all_posts.values()))
    week_start = date.fromisoformat(first["week_start"])
    week_end   = date.fromisoformat(first["week_end"])
    week_label = f"{week_end.year}-W{_week_number(week_start):02d}"

    output_dir = Path(args.output_dir or f"output/{week_label}")
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Documenten opslaan in: {output_dir.resolve()}\n")

    succeeded = 0
    failed    = []

    for i, (klant_id, data) in enumerate(all_posts.items(), 1):
        client    = data["client"]
        posts     = data["posts"]
        naam      = client["bedrijfsnaam"]
        week_nr   = _week_number(week_start)

        filename  = f"{naam} - Social Media Week {week_nr} {week_end.year}.docx"
        filepath  = output_dir / filename

        print(f"  [{i}/{len(all_posts)}] {naam}...", end=" ", flush=True)

        try:
            docx_bytes = _build_docx(naam, posts, week_start, week_end)
            filepath.write_bytes(docx_bytes)
            succeeded += 1
            print(f"✓  {filename}")
        except Exception as e:
            print(f"✗  {e}")
            failed.append((klant_id, str(e)))

    print(f"\n{'='*40}")
    print(f"✓ Opgeslagen: {succeeded}/{len(all_posts)} documenten")
    if failed:
        print(f"✗ Mislukt ({len(failed)}):")
        for klant_id, reason in failed:
            print(f"  - {klant_id}: {reason}")
    print(f"\nMap openen:")
    print(f"  open \"{output_dir.resolve()}\"")


if __name__ == "__main__":
    main()
