"""
Maakt per klant een Word-document (.docx) aan en uploadt dit naar een SharePoint-map
via de Microsoft Graph API.

Authenticatie via client credentials (app-only) — geen gebruikerslogin nodig.

Vereiste kolom in Google Sheet: sharepoint_folder_path
  Relatief pad binnen de SharePoint-documentbibliotheek, bijv:
  Klanten/Antwan Tuinprojecten

Vereiste .env variabelen:
  MICROSOFT_CLIENT_ID
  MICROSOFT_CLIENT_SECRET
  MICROSOFT_TENANT_ID
  SHAREPOINT_SITE_URL     bijv. https://jouwbedrijf.sharepoint.com/sites/SocialMedia
  SHAREPOINT_DRIVE_NAME   bijv. Documenten (naam van de documentbibliotheek, default: Documents)

Usage:
    python systems/create_sharepoint_doc.py
    python systems/create_sharepoint_doc.py --posts-file intermediates/posts_2026-06-05.json
    python systems/create_sharepoint_doc.py --client-id DR-001
"""

import argparse
import io
import json
import os
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Optional

import msal
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv(override=True)

GRAPH_URL = "https://graph.microsoft.com/v1.0"
SCOPE = ["https://graph.microsoft.com/.default"]

MONTHS_NL = [
    "januari", "februari", "maart", "april", "mei", "juni",
    "juli", "augustus", "september", "oktober", "november", "december",
]

PLATFORM_COLORS = {
    "instagram": RGBColor(0xE1, 0x30, 0x6C),
    "linkedin":  RGBColor(0x00, 0x77, 0xB5),
    "facebook":  RGBColor(0x18, 0x77, 0xF2),
}

PLATFORM_LABELS = {
    "instagram": "Instagram",
    "linkedin":  "LinkedIn",
    "facebook":  "Facebook",
}


def _get_access_token() -> str:
    client_id     = os.getenv("MICROSOFT_CLIENT_ID")
    client_secret = os.getenv("MICROSOFT_CLIENT_SECRET")
    tenant_id     = os.getenv("MICROSOFT_TENANT_ID")

    for var in ("MICROSOFT_CLIENT_ID", "MICROSOFT_CLIENT_SECRET", "MICROSOFT_TENANT_ID"):
        if not os.getenv(var):
            raise EnvironmentError(f"{var} not set in .env")

    app = msal.ConfidentialClientApplication(
        client_id,
        authority=f"https://login.microsoftonline.com/{tenant_id}",
        client_credential=client_secret,
    )
    result = app.acquire_token_for_client(scopes=SCOPE)
    if "access_token" not in result:
        raise RuntimeError(f"Token ophalen mislukt: {result.get('error_description')}")
    return result["access_token"]


def _graph(token: str, method: str, path: str, **kwargs) -> requests.Response:
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    headers.update(kwargs.pop("headers", {}))
    r = requests.request(method, f"{GRAPH_URL}{path}", headers=headers, **kwargs)
    r.raise_for_status()
    return r


def _get_site_id(token: str, site_url: str) -> str:
    """Resolves a SharePoint site URL to a Graph site ID."""
    from urllib.parse import urlparse
    parsed = urlparse(site_url)
    hostname = parsed.netloc
    site_path = parsed.path.rstrip("/")
    r = _graph(token, "GET", f"/sites/{hostname}:{site_path}")
    return r.json()["id"]


def _get_drive_id(token: str, site_id: str, drive_name: str) -> str:
    """Finds the drive (document library) ID by name."""
    r = _graph(token, "GET", f"/sites/{site_id}/drives")
    drives = r.json().get("value", [])
    for drive in drives:
        if drive["name"].lower() == drive_name.lower():
            return drive["id"]
    # Fallback: return default drive
    if drives:
        return drives[0]["id"]
    raise RuntimeError(f"Geen documentbibliotheek gevonden met naam '{drive_name}'")


def _ensure_folder_path(token: str, drive_id: str, folder_path: str) -> str:
    """Creates folder path if it doesn't exist. Returns the final folder item ID."""
    parts = [p for p in folder_path.strip("/").split("/") if p]
    current_id = "root"

    for part in parts:
        try:
            r = _graph(token, "GET", f"/drives/{drive_id}/items/{current_id}/children")
            children = r.json().get("value", [])
            match = next((c for c in children if c["name"].lower() == part.lower()), None)
            if match:
                current_id = match["id"]
            else:
                # Create folder
                body = {"name": part, "folder": {}, "@microsoft.graph.conflictBehavior": "rename"}
                r = _graph(token, "POST", f"/drives/{drive_id}/items/{current_id}/children", json=body)
                current_id = r.json()["id"]
        except Exception as e:
            raise RuntimeError(f"Map '{part}' aanmaken mislukt: {e}")

    return current_id


def _format_date_nl(d: date) -> str:
    return f"{d.day} {MONTHS_NL[d.month - 1]}"


def _week_number(d: date) -> int:
    return d.isocalendar()[1]


def _build_docx(bedrijfsnaam: str, posts: dict, week_start: date, week_end: date) -> bytes:
    """Creates a Word document and returns it as bytes."""
    doc = Document()

    # Stel marges in
    for section in doc.sections:
        section.top_margin    = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin   = Pt(70)
        section.right_margin  = Pt(70)

    week_nr = _week_number(week_start)
    title_text = (
        f"{bedrijfsnaam} — Social Media"
        f" | Week {week_nr}"
        f" | {_format_date_nl(week_start)} – {_format_date_nl(week_end)} {week_end.year}"
    )

    # Titel
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitel
    sub = doc.add_paragraph(f"Gegenereerd op {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    for platform_key in ("instagram", "linkedin", "facebook"):
        platform_posts = posts.get(platform_key, [])
        if not platform_posts:
            continue

        label = PLATFORM_LABELS[platform_key]
        color = PLATFORM_COLORS[platform_key]
        count = len(platform_posts)

        # Platform-header
        heading = doc.add_heading(
            f"{label} — {count} post{'s' if count > 1 else ''}",
            level=2,
        )
        for run in heading.runs:
            run.font.color.rgb = color

        for post in platform_posts:
            dag     = post.get("dag", "")
            caption = post.get("caption", "")
            hashtags = post.get("hashtags", "")

            # Dag
            dag_p = doc.add_paragraph()
            dag_run = dag_p.add_run(f"📅 {dag}")
            dag_run.bold = True
            dag_run.font.size = Pt(11)

            # Caption
            doc.add_paragraph(caption)

            # Hashtags
            ht_p = doc.add_paragraph()
            ht_run = ht_p.add_run(hashtags)
            ht_run.font.color.rgb = color
            ht_run.font.size = Pt(10)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_doc(token: str, site_id: str, drive_id: str, client: dict,
               posts: dict, week_start: date, week_end: date) -> dict:
    bedrijfsnaam  = client["bedrijfsnaam"]
    folder_path   = client.get("sharepoint_folder_path", "").strip()

    if not folder_path:
        raise ValueError(f"Geen sharepoint_folder_path voor {bedrijfsnaam}")

    week_nr  = _week_number(week_start)
    filename = (
        f"{bedrijfsnaam} - Social Media Week {week_nr} {week_end.year}.docx"
    )

    docx_bytes = _build_docx(bedrijfsnaam, posts, week_start, week_end)
    folder_id  = _ensure_folder_path(token, drive_id, folder_path)

    # Upload
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    r = requests.put(
        f"{GRAPH_URL}/drives/{drive_id}/items/{folder_id}:/{filename}:/content",
        headers=headers,
        data=docx_bytes,
    )
    r.raise_for_status()
    file_data = r.json()

    return {
        "doc_url":  file_data.get("webUrl", ""),
        "doc_id":   file_data.get("id", ""),
        "filename": filename,
    }


def _latest_file(pattern: str) -> Optional[str]:
    files = sorted(Path("intermediates").glob(pattern), reverse=True)
    return str(files[0]) if files else None


def main():
    parser = argparse.ArgumentParser(description="Maak Word-documenten aan in SharePoint")
    parser.add_argument("--posts-file", help="Pad naar posts JSON")
    parser.add_argument("--client-id", help="Verwerk alleen deze klant")
    parser.add_argument("--output", help="Pad naar rapportbestand")
    args = parser.parse_args()

    site_url   = os.getenv("SHAREPOINT_SITE_URL")
    drive_name = os.getenv("SHAREPOINT_DRIVE_NAME", "Documents")
    if not site_url:
        raise EnvironmentError("SHAREPOINT_SITE_URL not set in .env")

    posts_file = args.posts_file or _latest_file("posts_*.json")
    if not posts_file or not Path(posts_file).exists():
        print("Geen posts-bestand gevonden. Voer eerst generate_weekly_posts.py uit.", file=sys.stderr)
        sys.exit(1)

    with open(posts_file, encoding="utf-8") as f:
        all_posts = json.load(f)

    if args.client_id:
        if args.client_id not in all_posts:
            print(f"Klant '{args.client_id}' niet in posts-bestand.", file=sys.stderr)
            sys.exit(1)
        all_posts = {args.client_id: all_posts[args.client_id]}

    print("Microsoft token ophalen...")
    token = _get_access_token()

    print(f"SharePoint site ophalen: {site_url}")
    site_id  = _get_site_id(token, site_url)
    drive_id = _get_drive_id(token, site_id, drive_name)
    print(f"Documentbibliotheek: {drive_name}")

    output_path = args.output or f"intermediates/report_{date.today()}.json"
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    report = {}
    if Path(output_path).exists():
        with open(output_path, encoding="utf-8") as f:
            report = json.load(f)
        print(f"Hervat: {len(report)} docs al aangemaakt.")

    to_process = {k: v for k, v in all_posts.items() if k not in report}
    print(f"{len(to_process)} Word-documenten aan te maken...\n")

    succeeded = 0
    failed    = []

    for i, (klant_id, data) in enumerate(to_process.items(), 1):
        client     = data["client"]
        posts      = data["posts"]
        week_start = date.fromisoformat(data["week_start"])
        week_end   = date.fromisoformat(data["week_end"])

        print(f"  [{i}/{len(to_process)}] {client['bedrijfsnaam']}...", end=" ", flush=True)

        try:
            result = create_doc(token, site_id, drive_id, client, posts, week_start, week_end)
            report[klant_id] = result
            succeeded += 1
            print(f"✓  {result['doc_url']}")
        except Exception as e:
            print(f"✗  {e}")
            failed.append((klant_id, str(e)))

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n{'='*40}")
    print(f"✓ Aangemaakt: {succeeded}/{len(to_process)} documenten")
    if failed:
        print(f"✗ Mislukt ({len(failed)}):")
        for klant_id, reason in failed:
            print(f"  - {klant_id}: {reason}")
    print(f"Rapport opgeslagen: {output_path}")


if __name__ == "__main__":
    main()
