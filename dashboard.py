"""
Social Media Agent Dashboard

Toont actieve klanten en de volgende geplande run.
Leest klantdata live uit Google Sheets.

Lokaal starten:
    streamlit run dashboard.py

Vereist in .env (of Streamlit Cloud secrets):
    GOOGLE_SERVICE_ACCOUNT_JSON
    GOOGLE_SHEETS_SPREADSHEET_ID
"""

import base64
import hashlib
import hmac
import io
import json
import os
import re
import time
import zipfile
from datetime import datetime, date, timedelta
from datetime import time as dtime
from urllib.parse import urlparse
from zoneinfo import ZoneInfo

import gspread
import streamlit as st
import streamlit.components.v1 as components
import bcrypt
from google.oauth2.service_account import Credentials as WriteCredentials
from google.oauth2.service_account import Credentials
from dotenv import load_dotenv

from systems import drive_upload
from systems import publish_scheduled_posts as pub

load_dotenv(override=True)

# Streamlit Cloud draait op UTC — voor de juiste runtijden en datums
# rekenen we expliciet om naar Nederlandse (Amsterdamse) tijd.
AMSTERDAM_TZ = ZoneInfo("Europe/Amsterdam")


def _now_ams() -> datetime:
    """Huidige tijd in Amsterdam (incl. automatische zomer-/wintertijd), tz-naïef
    gemaakt zodat het probleemloos te combineren is met de rest van de datetime-logica."""
    return datetime.now(AMSTERDAM_TZ).replace(tzinfo=None)

# ── Paginaconfiguratie ────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Top Socials · TopMediaGroep",
    page_icon="https://www.topmediagroep.nl/data/pam/public/logo/logo_topmediagroep_transparent.png",
    layout="wide",
)

# ── Helpers ───────────────────────────────────────────────────────────────────

PLATFORM_COLORS = {
    "instagram": "#E1306C",
    "linkedin":  "#0077B5",
    "facebook":  "#1877F2",
}

def _platform_badge(label: str, color: str) -> str:
    """Zelfvoorzienend platform-icoon (geen externe afbeelding nodig — rendert altijd,
    zodat icoon en tekst nooit door elkaar kunnen lopen door een mislukte image-load)."""
    return (
        f'<span style="display:inline-flex;align-items:center;justify-content:center;'
        f'width:20px;height:20px;min-width:20px;border-radius:6px;background:{color};'
        f'color:#fff;font-size:9px;font-weight:800;letter-spacing:.02em;'
        f'margin-right:7px;vertical-align:middle;line-height:1;">{label}</span>'
    )


PLATFORM_ICON_HTML = {
    "instagram": _platform_badge("IG", PLATFORM_COLORS["instagram"]),
    "linkedin":  _platform_badge("IN", PLATFORM_COLORS["linkedin"]),
    "facebook":  _platform_badge("FB", PLATFORM_COLORS["facebook"]),
}

PLATFORM_LABELS = {
    "instagram": "Instagram",
    "linkedin":  "LinkedIn",
    "facebook":  "Facebook",
}

SCOPES = ["https://www.googleapis.com/auth/spreadsheets.readonly"]

WRITE_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]


# ── Auth helpers (vroeg definiëren — worden vóór CSS/content aangeroepen) ────

def _auth_users() -> dict:
    """Geeft alle gebruikers uit st.secrets['auth']['users'] terug, of {} bij lokale dev."""
    try:
        return dict(st.secrets.get("auth", {}).get("users", {}))
    except Exception:
        return {}


def _verify_password(plain: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(plain.encode(), hashed.encode())
    except Exception:
        return False


def _current_user_role(username: str) -> str:
    """Geeft de rol van de ingelogde gebruiker: 'admin' of 'reviewer'."""
    try:
        return _auth_users()[username].get("role", "reviewer")
    except Exception:
        return "reviewer"


# ── Persistente login (12 uur via signed cookie) ─────────────────────────────
SESSION_HOURS = 12
AUTH_COOKIE_NAME = "ts_auth_token"


def _cookie_secret() -> str:
    try:
        return str(st.secrets.get("auth", {}).get("cookie_secret", "top-socials-dev-secret"))
    except Exception:
        return "top-socials-dev-secret"


def _sign(payload: str) -> str:
    return hmac.new(_cookie_secret().encode(), payload.encode(), hashlib.sha256).hexdigest()


def _make_session_token(username: str) -> str:
    """Maakt een ondertekend token dat 12 uur geldig is."""
    expiry = int(time.time()) + SESSION_HOURS * 3600
    payload = f"{username}|{expiry}"
    token = f"{payload}|{_sign(payload)}"
    return base64.urlsafe_b64encode(token.encode()).decode()


def _verify_session_token(token: str):
    """Geeft de gebruikersnaam terug als het token geldig en niet verlopen is, anders None."""
    try:
        raw = base64.urlsafe_b64decode(token.encode()).decode()
        username, expiry, sig = raw.split("|")
        if not hmac.compare_digest(_sign(f"{username}|{expiry}"), sig):
            return None
        if int(expiry) < time.time():
            return None
        return username
    except Exception:
        return None


def _sa_info_json() -> str | None:
    b64 = os.getenv("GOOGLE_SERVICE_ACCOUNT_B64") or st.secrets.get("GOOGLE_SERVICE_ACCOUNT_B64")
    if not b64:
        p1 = st.secrets.get("GOOGLE_SA_B64_1", "")
        p2 = st.secrets.get("GOOGLE_SA_B64_2", "")
        b64 = p1 + p2 if p1 and p2 else None
    if b64:
        return json.dumps(json.loads(base64.b64decode(b64).decode()))
    sa = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    return sa if sa else None


@st.cache_data(ttl=30)
def load_medewerker_assignments(spreadsheet_id: str, sa_info_json: str) -> dict:
    """Laadt klant-toewijzingen uit de 'Medewerkers'-tab van de Google Sheet."""
    try:
        sa_info = json.loads(sa_info_json)
        creds   = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        gc      = gspread.authorize(creds)
        sh      = gc.open_by_key(spreadsheet_id)
        try:
            ws = sh.worksheet("Medewerkers")
        except gspread.WorksheetNotFound:
            return {}
        rows   = ws.get_all_records(default_blank="")
        result = {}
        for row in rows:
            uname = str(row.get("gebruiker", "")).strip().lower()
            raw   = str(row.get("klanten",   "")).strip()
            if not uname:
                continue
            result[uname] = "ALL" if raw.upper() == "ALL" else [
                k.strip() for k in raw.split(",") if k.strip()
            ]
        return result
    except Exception:
        return {}


def save_medewerker_assignments(spreadsheet_id: str, sa_info_json: str,
                                assignments: dict) -> None:
    """Schrijft klant-toewijzingen terug naar de 'Medewerkers'-tab."""
    sa_info = json.loads(sa_info_json)
    creds   = WriteCredentials.from_service_account_info(sa_info, scopes=WRITE_SCOPES)
    gc      = gspread.authorize(creds)
    sh      = gc.open_by_key(spreadsheet_id)
    try:
        ws = sh.worksheet("Medewerkers")
        ws.clear()
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet("Medewerkers", rows=100, cols=2)
    rows = [["gebruiker", "klanten"]]
    for uname, clients in assignments.items():
        rows.append([uname, "ALL" if clients == "ALL" else ", ".join(clients)])
    ws.update(rows, value_input_option="RAW")
    ws.freeze(rows=1)
    ws.format("A1:B1", {"textFormat": {"bold": True}})
    load_medewerker_assignments.clear()


def _next_wednesday_23() -> datetime:
    """Geeft de eerstvolgende woensdag om 23:00 Amsterdam-tijd terug (tz-naïeve datetime in Amsterdamse tijd)."""
    now = _now_ams()
    days_ahead = (2 - now.weekday()) % 7  # 2 = woensdag
    if days_ahead == 0 and now.hour >= 23:
        days_ahead = 7
    next_run = now.replace(hour=23, minute=0, second=0, microsecond=0) + timedelta(days=days_ahead)
    return next_run


def _format_countdown(delta: timedelta) -> str:
    total_seconds = int(delta.total_seconds())
    days    = total_seconds // 86400
    hours   = (total_seconds % 86400) // 3600
    minutes = (total_seconds % 3600) // 60
    if days > 0:
        return f"{days}d {hours}u {minutes}m"
    return f"{hours}u {minutes}m"


def get_profile_image(website_url: str) -> str:
    """Geeft Google favicon-URL terug — geen scraping, altijd instant."""
    if not website_url:
        return ""
    domain = urlparse(website_url).netloc
    return f"https://www.google.com/s2/favicons?domain={domain}&sz=128" if domain else ""


def _client_follower_counts(client: dict) -> dict:
    """Leest opgeslagen volgersaantallen uit de Google Sheet — geen live scraping."""
    return {
        "instagram": client.get("instagram_volgers", "") or "—",
        "linkedin":  client.get("linkedin_volgers",  "") or "—",
        "facebook":  client.get("facebook_volgers",  "") or "—",
    }


@st.cache_data(ttl=300)  # Cache 5 minuten
def load_clients() -> list[dict]:
    # Lees spreadsheet ID
    spreadsheet_id = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID")
    if not spreadsheet_id:
        st.error("GOOGLE_SHEETS_SPREADSHEET_ID ontbreekt in secrets.")
        return []

    # Lees service account — base64 string heeft voorkeur (geen TOML-escapingproblemen)
    try:
        b64 = (os.getenv("GOOGLE_SERVICE_ACCOUNT_B64")
               or st.secrets.get("GOOGLE_SERVICE_ACCOUNT_B64"))
        if not b64:
            part1 = st.secrets.get("GOOGLE_SA_B64_1", "")
            part2 = st.secrets.get("GOOGLE_SA_B64_2", "")
            if part1 and part2:
                b64 = part1 + part2
        if b64:
            sa_info = json.loads(base64.b64decode(b64).decode())
        else:
            sa_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
            if not sa_json:
                st.error("GOOGLE_SERVICE_ACCOUNT_B64 ontbreekt in secrets.")
                return []
            sa_info = json.loads(sa_json)
    except Exception as e:
        st.error(f"Fout bij laden service account: {e}")
        return []

    try:
        creds  = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        client = gspread.authorize(creds)
        sheet  = client.open_by_key(spreadsheet_id).sheet1
        rows   = sheet.get_all_records(default_blank="")
    except Exception as e:
        st.error(f"Fout bij verbinden met Google Sheets: {e}")
        return []

    active = []
    for row in rows:
        if str(row.get("actief", "")).strip().upper() != "TRUE":
            continue
        for platform in ("instagram", "linkedin", "facebook"):
            key = f"{platform}_posts_pw"
            try:
                row[key] = int(row[key])
            except (ValueError, TypeError):
                row[key] = 0
        active.append(row)

    return active


def _total_posts_pw(client: dict) -> int:
    return sum(client.get(f"{p}_posts_pw", 0) for p in ("instagram", "linkedin", "facebook"))


STATS_SHEET_NAME = "Statistieken"

STATS_NUMERIC_COLUMNS = [
    "instagram_volgers", "instagram_bereik_7d", "instagram_impressies_7d", "instagram_profielbezoeken_7d",
    "facebook_volgers", "facebook_bereik_7d", "facebook_impressies_7d", "facebook_engagement_7d",
]


@st.cache_data(ttl=300)  # Cache 5 minuten
def load_statistics() -> list[dict]:
    """Leest het tabblad 'Statistieken' (Meta-koppeling). Geeft [] terug als het
    tabblad nog niet bestaat (bijv. nog geen klanten gekoppeld)."""
    spreadsheet_id = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID")
    if not spreadsheet_id:
        return []

    try:
        b64 = (os.getenv("GOOGLE_SERVICE_ACCOUNT_B64")
               or st.secrets.get("GOOGLE_SERVICE_ACCOUNT_B64"))
        if not b64:
            part1 = st.secrets.get("GOOGLE_SA_B64_1", "")
            part2 = st.secrets.get("GOOGLE_SA_B64_2", "")
            if part1 and part2:
                b64 = part1 + part2
        if b64:
            sa_info = json.loads(base64.b64decode(b64).decode())
        else:
            sa_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
            if not sa_json:
                return []
            sa_info = json.loads(sa_json)
    except Exception:
        return []

    try:
        creds  = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        client = gspread.authorize(creds)
        sheet  = client.open_by_key(spreadsheet_id)
        try:
            ws = sheet.worksheet(STATS_SHEET_NAME)
        except gspread.exceptions.WorksheetNotFound:
            return []
        rows = ws.get_all_records(default_blank="")
    except Exception:
        return []

    for row in rows:
        for col in STATS_NUMERIC_COLUMNS:
            try:
                row[col] = float(row[col]) if row.get(col, "") != "" else None
            except (ValueError, TypeError):
                row[col] = None

    return rows


def _latest_stats_per_client(stats: list[dict]) -> dict:
    """Geeft per klant_id de meest recente meting terug."""
    latest = {}
    for row in stats:
        klant_id = row.get("klant_id", "")
        if not klant_id:
            continue
        if klant_id not in latest or row.get("datum", "") >= latest[klant_id].get("datum", ""):
            latest[klant_id] = row
    return latest


def _format_stat(value) -> str:
    if value is None:
        return "—"
    value = int(value)
    if value >= 1000:
        return f"{value / 1000:.1f}k".replace(".0k", "k")
    return str(value)


def _stat_delta(history: list[dict], col: str):
    """Verschil t.o.v. de meting daarvoor, of None als er geen vorige meting is."""
    values = [r[col] for r in history if r.get(col) is not None]
    if len(values) < 2:
        return None
    diff = values[-1] - values[-2]
    if diff == 0:
        return None
    sign = "+" if diff > 0 else ""
    return f"{sign}{int(diff)} t.o.v. vorige meting"


def _followers_growth(history: list[dict], col: str, days: int = 7):
    """Geeft (verschil, percentage) volgersgroei t.o.v. ~`days` dagen geleden.

    Gebruikt de oudste meting binnen de laatste `days` dagen als basislijn.
    Geeft None terug als er geen meting met deze kolom beschikbaar is.
    """
    values = []
    for r in history:
        if r.get(col) is None or not r.get("datum"):
            continue
        try:
            d = datetime.strptime(r["datum"], "%Y-%m-%d").date()
        except ValueError:
            continue
        values.append((d, r[col]))
    if len(values) < 2:
        return None

    latest_date, latest_val = values[-1]
    cutoff = latest_date - timedelta(days=days)
    baseline = values[0]
    for d, val in values:
        if d <= cutoff:
            baseline = (d, val)
        else:
            break

    base_val = baseline[1]
    diff = latest_val - base_val
    pct = (diff / base_val * 100) if base_val else None
    return diff, pct


# ── Posts-statistieken & demografie (Meta-koppeling, vervolg) ────────────────

POSTS_SHEET_NAME = "Statistieken_Posts"
DEMO_SHEET_NAME  = "Demografie"

DAGEN_NL = ["Maandag", "Dinsdag", "Woensdag", "Donderdag", "Vrijdag", "Zaterdag", "Zondag"]


def _to_float(value):
    """Parseert getallen die door Google Sheets als '0,0625' (NL-notatie) zijn opgeslagen."""
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    try:
        return float(str(value).replace(",", "."))
    except (ValueError, TypeError):
        return None


@st.cache_data(ttl=300)
def load_post_insights() -> list[dict]:
    """Leest het tabblad 'Statistieken_Posts'. Geeft [] terug als het nog niet bestaat."""
    spreadsheet_id = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID")
    if not spreadsheet_id:
        return []
    try:
        b64 = (os.getenv("GOOGLE_SERVICE_ACCOUNT_B64") or st.secrets.get("GOOGLE_SERVICE_ACCOUNT_B64"))
        if not b64:
            part1 = st.secrets.get("GOOGLE_SA_B64_1", "")
            part2 = st.secrets.get("GOOGLE_SA_B64_2", "")
            if part1 and part2:
                b64 = part1 + part2
        if b64:
            sa_info = json.loads(base64.b64decode(b64).decode())
        else:
            sa_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
            if not sa_json:
                return []
            sa_info = json.loads(sa_json)
    except Exception:
        return []

    try:
        creds  = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        client = gspread.authorize(creds)
        sheet  = client.open_by_key(spreadsheet_id)
        try:
            ws = sheet.worksheet(POSTS_SHEET_NAME)
        except gspread.exceptions.WorksheetNotFound:
            return []
        rows = ws.get_all_records(default_blank="")
    except Exception:
        return []

    for row in rows:
        row["bereik"] = _to_float(row.get("bereik"))
        row["interacties"] = _to_float(row.get("interacties"))
        row["engagement_rate"] = _to_float(row.get("engagement_rate"))

    return rows


@st.cache_data(ttl=300)
def load_demographics() -> list[dict]:
    """Leest het tabblad 'Demografie'. Geeft [] terug als het nog niet bestaat."""
    spreadsheet_id = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID")
    if not spreadsheet_id:
        return []
    try:
        b64 = (os.getenv("GOOGLE_SERVICE_ACCOUNT_B64") or st.secrets.get("GOOGLE_SERVICE_ACCOUNT_B64"))
        if not b64:
            part1 = st.secrets.get("GOOGLE_SA_B64_1", "")
            part2 = st.secrets.get("GOOGLE_SA_B64_2", "")
            if part1 and part2:
                b64 = part1 + part2
        if b64:
            sa_info = json.loads(base64.b64decode(b64).decode())
        else:
            sa_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
            if not sa_json:
                return []
            sa_info = json.loads(sa_json)
    except Exception:
        return []

    try:
        creds  = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        client = gspread.authorize(creds)
        sheet  = client.open_by_key(spreadsheet_id)
        try:
            ws = sheet.worksheet(DEMO_SHEET_NAME)
        except gspread.exceptions.WorksheetNotFound:
            return []
        rows = ws.get_all_records(default_blank="")
    except Exception:
        return []

    for row in rows:
        row["aantal"] = _to_float(row.get("aantal")) or 0

    return rows


def _latest_demo_per_client(demo: list[dict], klant_id: str) -> dict[str, list[dict]]:
    """Geeft de meest recente demografie-meting per dimensie voor een klant."""
    rows = [r for r in demo if r.get("klant_id") == klant_id]
    if not rows:
        return {}
    by_dim = {}
    for r in rows:
        by_dim.setdefault(r["dimensie"], []).append(r)
    out = {}
    for dim, items in by_dim.items():
        latest_date = max(r["datum"] for r in items)
        out[dim] = [r for r in items if r["datum"] == latest_date]
    return out


@st.cache_data(ttl=3600)
def generate_ai_summary(klant_id: str, bedrijfsnaam: str, latest: dict, previous: dict, top_posts: list[dict]) -> str:
    """Genereert een korte Nederlandse samenvatting van de recente prestaties via Claude."""
    api_key = os.getenv("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return ""

    import anthropic

    cijfers = []
    for label, key in [
        ("IG volgers", "instagram_volgers"), ("IG bereik (7d)", "instagram_bereik_7d"),
        ("FB volgers", "facebook_volgers"), ("FB bereik (7d)", "facebook_bereik_7d"),
        ("FB engagement (7d)", "facebook_engagement_7d"),
    ]:
        nu = latest.get(key)
        if nu is None:
            continue
        vorig = previous.get(key) if previous else None
        if vorig is not None:
            cijfers.append(f"{label}: {int(nu)} (vorige meting: {int(vorig)})")
        else:
            cijfers.append(f"{label}: {int(nu)}")

    posts_lines = []
    for p in top_posts[:3]:
        posts_lines.append(f"- ({p['platform']}) \"{p['caption_kort']}\" — engagement rate {p['engagement_rate']:.1%}" if p.get("engagement_rate") else f"- ({p['platform']}) \"{p['caption_kort']}\"")

    prompt = (
        f"Je bent een social media-analist. Schrijf een korte (max 3 zinnen), "
        f"vriendelijke Nederlandse samenvatting voor klant '{bedrijfsnaam}' op basis van "
        f"deze cijfers:\n\n" + "\n".join(cijfers) +
        ("\n\nBest presterende recente posts:\n" + "\n".join(posts_lines) if posts_lines else "") +
        "\n\nNoem concrete getallen, leg kort uit wat opvalt, en geef indien mogelijk één "
        "praktisch inzicht voor toekomstige content. Geen opsomming, gewoon lopende tekst."
    )

    try:
        ac = anthropic.Anthropic(api_key=api_key)
        message = ac.messages.create(
            model="claude-haiku-4-5",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        return message.content[0].text.strip()
    except Exception:
        return ""


# ── Merk & design system ──────────────────────────────────────────────────────
#
# "Top Socials" — het content-cockpit van TopMediaGroep.
# Eén consistente kleuren-, typografie- en componentenset zodat de tool oogt
# als een afgewerkt product i.p.v. een intern script.

BRAND = {
    "primary":      "#4F46E5",   # indigo — merkkleur
    "primary_dark": "#3730A3",
    "primary_soft": "#F0F0FF",
    "success":      "#34C759",   # Apple groen
    "warning":      "#FF9F0A",   # Apple oranje
    "danger":       "#FF3B30",   # Apple rood
    "ink":          "#1D1D1F",   # Apple near-black
    "ink_soft":     "#86868B",   # Apple secondary
    "ink_xsoft":    "#C7C7CC",   # Apple tertiary
    "line":         "rgba(0,0,0,.08)",
    "surface":      "#FFFFFF",
    "canvas":       "#F5F5F7",   # Apple signature gray
}

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:ital,opsz,wght@0,14..32,300;0,14..32,400;0,14..32,500;0,14..32,600;0,14..32,700;0,14..32,800;0,14..32,900&display=swap');

:root {{
    --p:     {BRAND['primary']};
    --pd:    {BRAND['primary_dark']};
    --ps:    {BRAND['primary_soft']};
    --ok:    {BRAND['success']};
    --warn:  {BRAND['warning']};
    --err:   {BRAND['danger']};
    --ink:   {BRAND['ink']};
    --ink2:  {BRAND['ink_soft']};
    --ink3:  {BRAND['ink_xsoft']};
    --line:  {BRAND['line']};
    --surf:  {BRAND['surface']};
    --bg:    {BRAND['canvas']};
    --r-card: 20px;
    --r-btn:  10px;
    --r-inp:  12px;
    --shadow: 0 2px 12px rgba(0,0,0,.07), 0 0 0 0.5px rgba(0,0,0,.04);
    --shadow-hover: 0 8px 32px rgba(0,0,0,.12), 0 0 0 0.5px rgba(0,0,0,.05);
}}

/* ── Font ── */
html, body, .stApp,
.stMarkdown, .stMarkdown p, .stMarkdown li,
.stMarkdown span:not([class*="material"]),
.stCaption, .stText, p, label,
.stButton button, .stDownloadButton button, .stLinkButton a,
input, textarea, select,
[data-testid="stMetricLabel"], [data-testid="stMetricValue"],
[data-testid="stWidgetLabel"] {{
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    -webkit-font-smoothing: antialiased;
}}
[data-testid*="Icon"], [class*="material-symbols"],
[class*="MaterialSymbols"], .stApp [class*="icon"] {{
    font-family: 'Material Symbols Rounded', 'Material Icons' !important;
}}

/* ── Canvas ── */
html, body, .stApp,
[data-testid="stAppViewContainer"], [data-testid="stMain"],
.main, .main > div {{ background: var(--bg) !important; }}
.block-container {{ padding: 2rem 2.5rem 5rem !important; max-width: 1240px !important; }}

/* ── Streamlit chrome verbergen ── */
#MainMenu, footer, header[data-testid="stHeader"] {{
    visibility: hidden; height: 0; overflow: hidden;
}}
[data-testid="stToolbar"], [data-testid="manage-app-button"],
.stDeployButton, [data-testid="stDecoration"], [data-testid="stStatusWidget"],
[class*="viewerBadge"], [data-testid="stAppViewerBadge"] {{
    display: none !important;
}}

/* ── Koppen ── */
h1, h2, h3, h4 {{
    font-family: 'Inter', sans-serif !important;
    font-weight: 700 !important;
    color: var(--ink) !important;
    letter-spacing: -0.03em !important;
}}

/* ── Divider ── */
hr {{ border-color: var(--line) !important; margin: 1.5rem 0 !important; }}

/* ── KPI-kaarten ── */
[data-testid="stMetric"] {{
    background: var(--surf);
    border: none;
    border-radius: var(--r-card);
    padding: 20px 24px;
    box-shadow: var(--shadow);
    height: 132px;
    box-sizing: border-box;
    display: flex; flex-direction: column; justify-content: center;
}}
[data-testid="stMetricLabel"] {{
    font-size: 11px !important; font-weight: 600 !important;
    color: var(--ink2) !important; text-transform: uppercase; letter-spacing: .08em;
}}
[data-testid="stMetricValue"] {{
    font-size: 26px !important; font-weight: 700 !important;
    color: var(--ink) !important; letter-spacing: -0.02em;
}}
[data-testid="stMetricDelta"] svg {{ display: none; }}
[data-testid="stMetricDelta"] > div {{
    font-size: 12px !important; color: var(--ink2) !important;
    font-weight: 500 !important;
}}

/* ── Tabs: gesegmenteerde pill-stijl ── */
.stTabs [data-baseweb="tab-list"] {{
    background: rgba(0,0,0,.05);
    border-radius: 14px;
    padding: 4px;
    gap: 2px;
    border: none !important;
}}
.stTabs [data-baseweb="tab"] {{
    border-radius: 11px;
    padding: 7px 18px;
    font-weight: 600; font-size: 13px;
    color: var(--ink2);
    background: transparent;
    border: none !important;
    transition: all .18s ease;
}}
.stTabs [aria-selected="true"] {{
    background: var(--surf) !important;
    color: var(--ink) !important;
    box-shadow: 0 1px 6px rgba(0,0,0,.14) !important;
    border: none !important;
}}
.stTabs [data-baseweb="tab-highlight"] {{ display: none; }}
.stTabs [data-baseweb="tab-border"]    {{ display: none; }}

/* ── Knoppen ── */
.stButton > button, .stDownloadButton > button,
.stButton button, .stDownloadButton button,
[data-testid^="stBaseButton"] {{
    border-radius: var(--r-btn) !important;
    font-weight: 600 !important; font-size: 13px !important;
    border: 1px solid var(--line) !important;
    background: var(--surf) !important;
    color: var(--ink) !important;
    transition: all .18s cubic-bezier(.4,0,.2,1);
    letter-spacing: -.01em;
}}
.stButton > button[kind="primary"],
[data-testid="stBaseButton-primary"] {{
    background: var(--p) !important;
    border-color: var(--p) !important;
    color: #fff !important;
    box-shadow: 0 2px 10px rgba(79,70,229,.35) !important;
}}
.stButton > button:hover:not(:disabled),
[data-testid^="stBaseButton"]:hover:not(:disabled) {{
    box-shadow: var(--shadow) !important;
    transform: translateY(-1px);
}}
.stButton > button[kind="primary"]:hover:not(:disabled),
[data-testid="stBaseButton-primary"]:hover:not(:disabled) {{
    background: var(--pd) !important;
    border-color: var(--pd) !important;
}}

/* ── Inputs ── */
.stTextInput input, .stTextArea textarea,
.stSelectbox div[data-baseweb="select"] > div {{
    border-radius: var(--r-inp) !important;
    border: 1px solid var(--line) !important;
    background: var(--surf) !important;
    font-size: 14px !important;
    color: var(--ink) !important;
}}
.stTextInput input:focus, .stTextArea textarea:focus {{
    border-color: var(--p) !important;
    box-shadow: 0 0 0 3px rgba(79,70,229,.12) !important;
}}

/* ── Expanders ── */
[data-testid="stExpander"] {{
    background: var(--surf) !important;
    border: none !important;
    border-radius: 16px !important;
    box-shadow: var(--shadow) !important;
    overflow: hidden;
}}
[data-testid="stExpander"] summary {{
    font-weight: 600 !important; font-size: 13px !important;
    color: var(--ink2) !important;
    padding: 12px 16px !important;
}}

/* ── Sidebar ── */
[data-testid="stSidebar"] {{
    background: var(--surf) !important;
    border-right: 1px solid var(--line) !important;
}}
[data-testid="stSidebar"] .block-container {{
    padding: 1.5rem 1.2rem !important;
}}

/* ── Merk-header ── */
.ts-header {{
    display: flex; align-items: center;
    justify-content: space-between;
    padding: 18px 28px;
    margin-bottom: 22px;
    background: var(--surf);
    border-radius: var(--r-card);
    box-shadow: var(--shadow);
    position: relative;
    overflow: hidden;
}}
.ts-header::before {{
    content: "";
    position: absolute; left: 0; top: 0; bottom: 0; width: 5px;
    background: linear-gradient(180deg, var(--p), var(--pd));
}}
.ts-logo-wrap {{
    display: flex; align-items: center; justify-content: center;
    width: 52px; height: 52px; border-radius: 14px;
    background: var(--ps); flex-shrink: 0;
}}
.ts-logo {{
    height: 34px; width: auto;
    object-fit: contain; display: block;
}}
.ts-name {{
    font-size: 21px; font-weight: 800;
    color: var(--ink); letter-spacing: -0.03em;
    line-height: 1.2;
}}
.ts-brandtag {{
    display: inline-block; margin-left: 8px;
    font-size: 10px; font-weight: 700; letter-spacing: .08em;
    color: var(--p); background: var(--ps);
    padding: 2px 8px; border-radius: 99px; vertical-align: middle;
    text-transform: uppercase;
}}
.ts-sub {{
    font-size: 12px; font-weight: 400;
    color: var(--ink2); margin-top: 3px; letter-spacing: .01em;
}}
.ts-live {{
    display: inline-flex; align-items: center; gap: 6px;
    background: rgba(52,199,89,.12); color: #1a7f37;
    font-size: 11px; font-weight: 700; letter-spacing: .06em;
    padding: 6px 14px; border-radius: 99px;
}}
.ts-live-dot {{
    width: 6px; height: 6px; border-radius: 99px;
    background: var(--ok);
    box-shadow: 0 0 0 3px rgba(52,199,89,.25);
    animation: pulse 2s ease infinite;
}}
@keyframes pulse {{
    0%, 100% {{ opacity:1; }} 50% {{ opacity:.4; }}
}}

/* ── KPI-kaarten: gekleurde accentstrip per kaart ── */
[data-testid="stHorizontalBlock"]:has([data-testid="stMetric"]) > div:nth-of-type(1) [data-testid="stMetric"] {{
    border-top: 3px solid var(--p) !important;
}}
[data-testid="stHorizontalBlock"]:has([data-testid="stMetric"]) > div:nth-of-type(2) [data-testid="stMetric"] {{
    border-top: 3px solid var(--ok) !important;
}}
[data-testid="stHorizontalBlock"]:has([data-testid="stMetric"]) > div:nth-of-type(3) [data-testid="stMetric"] {{
    border-top: 3px solid var(--warn) !important;
}}

/* ── Klant-tegels ── */
/* Container met key="tile-..." wordt de hele kaart (tegel + details samen) */
div[class*="st-key-tile-"] {{
    background: var(--surf);
    border-radius: var(--r-card);
    box-shadow: var(--shadow);
    overflow: hidden;
    margin-bottom: 28px;
    transition: box-shadow .2s ease, transform .2s ease;
}}
div[class*="st-key-tile-"]:hover {{
    box-shadow: var(--shadow-hover);
    transform: translateY(-2px);
}}
div[class*="st-key-tile-"] [data-testid="stVerticalBlock"] {{
    gap: 0 !important;
}}
div[class*="st-key-tile-"] [data-testid="stExpander"],
div[class*="st-key-tile-"] [data-testid="stExpander"] > details,
div[class*="st-key-tile-"] [data-testid="stExpander"] summary,
div[class*="st-key-tile-"] [data-testid="stExpander"] > div {{
    background: transparent !important;
    border: none !important;
    border-radius: 0 !important;
    box-shadow: none !important;
    margin: 0 !important;
}}
.ts-tile {{
    padding: 26px 26px 22px;
    cursor: default;
    height: 196px;
    box-sizing: border-box;
    display: flex; flex-direction: column; justify-content: space-between;
}}
.ts-tile-top {{
    display: flex; align-items: center; gap: 14px;
}}
.ts-tile-fav {{
    width: 46px; height: 46px; border-radius: 12px;
    background: var(--ps); flex-shrink: 0;
    object-fit: contain;
    border: 0.5px solid rgba(0,0,0,.06);
}}
.ts-tile-name {{
    font-size: 16px; font-weight: 700;
    color: var(--ink); letter-spacing: -0.02em;
    line-height: 1.35;
    display: -webkit-box; -webkit-line-clamp: 2;
    -webkit-box-orient: vertical; overflow: hidden;
}}
.ts-tile-badges {{
    display: flex; flex-wrap: wrap; gap: 8px; margin-top: 16px;
}}
.ts-tile-foot {{
    font-size: 12px; font-weight: 500;
    color: var(--ink3); letter-spacing: .01em;
    margin-top: 14px;
}}

/* ── Login-form: verberg Streamlit's eigen form-wrapper ── */
div[data-testid="stForm"] {{
    background: var(--surf) !important;
    border: none !important;
    border-radius: 22px !important;
    padding: 32px 28px 28px !important;
    box-shadow: 0 4px 30px rgba(0,0,0,.09), 0 0 0 0.5px rgba(0,0,0,.04) !important;
    max-width: 380px;
    margin: 0 auto !important;
}}
/* Vertaal Streamlit's "Press Enter to submit form"-hint naar Nederlands */
div[data-testid="stForm"] [data-testid="InputInstructions"] {{
    font-size: 0 !important;
}}
div[data-testid="stForm"] [data-testid="InputInstructions"]::after {{
    content: "Druk op Enter om te bevestigen";
    font-size: 12px;
    color: var(--ink3);
}}

/* ── iframe (Mail studio btn) ── */
iframe[title="components.html"], iframe[data-testid="stIFrame"] {{
    display: block; border: none !important;
    margin: 0 !important; vertical-align: top;
}}

/* ── Progress bar ── */
.ts-progress {{
    height: 4px; border-radius: 99px;
    background: rgba(0,0,0,.07); overflow: hidden; margin: 6px 0 14px;
}}
.ts-progress-fill {{
    height: 4px; border-radius: 99px;
    background: var(--ok); transition: width .4s ease;
}}

/* ── Statistieken-tab: kaarten ── */
div[class*="st-key-statcard-"] {{
    background: var(--surf);
    border-radius: var(--r-card);
    padding: 22px 24px;
    box-shadow: var(--shadow);
    margin-bottom: 18px;
}}
div[class*="st-key-statcard-"] [data-testid="stVerticalBlock"] {{
    gap: 0.4rem !important;
}}
/* Verberg de zoom/fullscreen-toolbar op grafieken in de Statistieken-tab */
div[class*="st-key-statcard-"] [data-testid="stElementToolbar"] {{
    display: none !important;
}}
.ts-stat-title {{
    font-size: 13px; font-weight: 700; color: var(--ink);
    margin: 0 0 2px; letter-spacing: -0.01em;
}}
.ts-stat-caption {{
    font-size: 12px; color: var(--ink2); margin: -2px 0 10px;
}}
.ts-post-card {{
    display: flex; align-items: center; gap: 14px;
    padding: 12px 14px; border-radius: 12px;
    background: var(--bg); border-left: 3px solid var(--p);
    margin-bottom: 8px;
}}
.ts-post-rank {{
    font-size: 18px; font-weight: 800; color: var(--ink3);
    width: 24px; text-align: center; flex-shrink: 0;
}}
.ts-post-thumb {{
    width: 56px; height: 56px; border-radius: 10px; object-fit: cover;
    flex-shrink: 0; background: var(--bg);
}}
.ts-post-thumb-empty {{
    display: flex; align-items: center; justify-content: center;
    font-size: 22px; color: var(--ink3);
}}
.ts-post-body {{ flex: 1; min-width: 0; }}
.ts-post-meta {{
    font-size: 11px; color: var(--ink2); margin-bottom: 2px;
    text-transform: uppercase; letter-spacing: .04em; font-weight: 600;
}}
.ts-post-caption {{
    font-size: 13px; color: var(--ink); line-height: 1.4;
    display: -webkit-box; -webkit-line-clamp: 2;
    -webkit-box-orient: vertical; overflow: hidden;
}}
.ts-post-stats {{
    display: flex; gap: 18px; flex-shrink: 0; text-align: right;
}}
.ts-post-stat {{ min-width: 56px; }}
.ts-post-stat-value {{ font-size: 15px; font-weight: 700; color: var(--ink); }}
.ts-post-stat-label {{
    font-size: 10px; color: var(--ink2); text-transform: uppercase; letter-spacing: .06em;
}}
.ts-post-link {{
    font-size: 16px; flex-shrink: 0; text-decoration: none;
}}
</style>
""", unsafe_allow_html=True)

# ── Favicon: TopMediaGroep-logo ───────────────────────────────────────────────
st.markdown(
    '<link rel="icon" type="image/png" href="https://www.topmediagroep.nl'
    '/data/pam/public/logo/logo_topmediagroep_transparent.png">',
    unsafe_allow_html=True,
)

# ── Authenticatie ────────────────────────────────────────────────────────────
_users = _auth_users()
_auth_active = bool(_users)

if not _auth_active:
    # Geen secrets → lokale dev, doorgaan als admin
    _logged_in_user = "dev"
    _logged_in_name = "Developer (lokaal)"
    _logged_in_role = "admin"
else:
    # ── Sessie herstellen via URL-token (blijft 12 uur geldig, ook na refresh) ──
    if not st.session_state.get("_ts_logged_in"):
        _token = st.query_params.get(AUTH_COOKIE_NAME)
        if _token:
            _restored_user = _verify_session_token(_token)
            if _restored_user and _restored_user in _users:
                st.session_state["_ts_logged_in"] = True
                st.session_state["_ts_username"]  = _restored_user
                st.session_state["_ts_name"]      = _users[_restored_user].get("name", _restored_user)
                st.session_state["_ts_role"]      = _users[_restored_user].get("role", "reviewer")
            else:
                # Ongeldig/verlopen token → verwijderen uit de URL
                del st.query_params[AUTH_COOKIE_NAME]

    # ── Login-check ───────────────────────────────────────────────────────────
    if not st.session_state.get("_ts_logged_in"):
        st.markdown(
            '<div style="max-width:380px;margin:100px auto 0;">'
            '<div style="text-align:center;margin-bottom:36px;">'
            '<img src="https://www.topmediagroep.nl/data/pam/public/logo/logo_topmediagroep_transparent.png" '
            'style="height:40px;width:auto;object-fit:contain;margin-bottom:20px;" alt="TopMediaGroep">'
            '<div style="font-size:22px;font-weight:800;color:#1D1D1F;letter-spacing:-.04em;'
            'font-family:Inter,-apple-system,sans-serif;">Top Socials</div>'
            '<div style="font-size:13px;color:#86868B;margin-top:5px;font-weight:400;">'
            'Content &amp; goedkeuring platform</div>'
            '</div>'
            '</div>',
            unsafe_allow_html=True,
        )

        with st.form("login_form", clear_on_submit=False, border=True):
            st.markdown('<p style="font-weight:700;font-size:17px;margin:0 0 16px;">Inloggen</p>',
                        unsafe_allow_html=True)
            username_input = st.text_input("E-mailadres", placeholder="naam@topmediagroep.nl")
            password_input = st.text_input("Wachtwoord", type="password")
            submitted = st.form_submit_button("Inloggen", use_container_width=True, type="primary")

        if submitted:
            user_key = username_input.strip().lower()
            user_data = _users.get(user_key) or _users.get(username_input.strip())
            if user_data and _verify_password(password_input, user_data.get("password", "")):
                st.session_state["_ts_logged_in"]   = True
                st.session_state["_ts_username"]    = user_key
                st.session_state["_ts_name"]        = user_data.get("name", user_key)
                st.session_state["_ts_role"]        = user_data.get("role", "reviewer")
                st.query_params[AUTH_COOKIE_NAME] = _make_session_token(user_key)
                st.rerun()
            else:
                st.error("❌ Onjuist e-mailadres of wachtwoord.")

        st.stop()

    _logged_in_user = st.session_state.get("_ts_username", "")
    _logged_in_name = st.session_state.get("_ts_name",     _logged_in_user)
    _logged_in_role = st.session_state.get("_ts_role",     "reviewer")

# ── Sidebar: gebruikersprofiel + uitlogknop ───────────────────────────────────
with st.sidebar:
    st.markdown(
        f'<div style="padding:8px 0 16px;">'
        f'<div style="font-size:14px;font-weight:700;color:#1D1D1F;letter-spacing:-.01em;">'
        f'{_logged_in_name}</div>'
        f'<div style="font-size:11px;font-weight:500;color:#86868B;margin-top:3px;'
        f'text-transform:uppercase;letter-spacing:.06em;">'
        f'{"Admin" if _logged_in_role == "admin" else "Medewerker"}'
        f'</div></div>',
        unsafe_allow_html=True,
    )
    if _auth_active:
        if st.button("Uitloggen", use_container_width=True):
            for k in ["_ts_logged_in", "_ts_username", "_ts_name", "_ts_role"]:
                st.session_state.pop(k, None)
            st.query_params.pop(AUTH_COOKIE_NAME, None)
            st.rerun()

# ── Merk-header ───────────────────────────────────────────────────────────────
st.markdown(
    '<div class="ts-header">'
    '  <div style="display:flex;align-items:center;gap:16px;">'
    '    <div class="ts-logo-wrap">'
    '      <img src="https://www.topmediagroep.nl/data/pam/public/logo/logo_topmediagroep_transparent.png" class="ts-logo" alt="TopMediaGroep">'
    '    </div>'
    '    <div>'
    '      <div class="ts-name">Top Socials<span class="ts-brandtag">TopMediaGroep</span></div>'
    '      <div class="ts-sub">Content &amp; goedkeuring platform</div>'
    '    </div>'
    '  </div>'
    '  <div class="ts-live"><div class="ts-live-dot"></div>LIVE</div>'
    '</div>',
    unsafe_allow_html=True,
)

# Volgende run
col1, col2, col3 = st.columns(3)

next_run = _next_wednesday_23()
delta    = next_run - _now_ams()

with col1:
    st.metric(
        label="Volgende run",
        value=next_run.strftime("woensdag %d %b om 23:00"),
        delta=f"over {_format_countdown(delta)}",
        delta_color="off",
    )

clients = load_clients()

# ── Filter klanten op basis van rol en medewerker-toewijzingen ───────────────
_all_clients = clients  # bewaar volledig voor admin-tab
if _logged_in_role != "admin":
    _sid_for_filter  = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID", "")
    _saj_for_filter  = _sa_info_json()
    if _sid_for_filter and _saj_for_filter:
        _assignments     = load_medewerker_assignments(_sid_for_filter, _saj_for_filter)
        _allowed_names   = _assignments.get(_logged_in_user, [])
        if _allowed_names != "ALL":
            clients = [c for c in clients if c.get("bedrijfsnaam", "") in _allowed_names]

with col2:
    st.metric("Actieve klanten", len(clients))

with col3:
    total_posts = sum(_total_posts_pw(c) for c in clients)
    st.metric("Posts per week", total_posts)

st.divider()

if _logged_in_role == "admin":
    tab_klanten, tab_goedkeuring, tab_planning, tab_statistieken, tab_team = st.tabs(
        ["📋 Klanten", "✅ Goedkeuring", "📅 Planning", "📊 Statistieken", "👥 Team"]
    )
else:
    tab_klanten, tab_goedkeuring, tab_planning, tab_statistieken = st.tabs(
        ["📋 Klanten", "✅ Goedkeuring", "📅 Planning", "📊 Statistieken"]
    )
    tab_team = None

with tab_klanten:
    if not clients:
        st.warning("Geen klanten gevonden. Controleer de Google Sheet en credentials.")
    else:
        # ── Zoekfilter ────────────────────────────────────────────────────────
        search = st.text_input("🔍 Zoek op klantnaam", placeholder="Typ een naam...",
                               label_visibility="collapsed")
        if search:
            clients = [c for c in clients if search.lower() in c.get("bedrijfsnaam", "").lower()]

        st.markdown(
            f'<p style="font-size:13px;color:#6B7280;margin:0 0 16px;">'
            f'{len(clients)} klant{"en" if len(clients) != 1 else ""} actief</p>',
            unsafe_allow_html=True,
        )

        # ── Tegeltjesraster: 3 per rij, ruim opgezet ─────────────────────────────
        COLS = 3
        for row_start in range(0, len(clients), COLS):
            row_clients = clients[row_start : row_start + COLS]
            cols = st.columns(COLS, gap="large")

            for col_idx, (col, client) in enumerate(zip(cols, row_clients)):
                total     = _total_posts_pw(client)
                img_url   = get_profile_image(client.get("website_url", ""))
                followers = _client_follower_counts(client)

                # Platform-badges (compact)
                badge_parts = []
                for platform, label in PLATFORM_LABELS.items():
                    count = client.get(f"{platform}_posts_pw", 0)
                    if count:
                        c = PLATFORM_COLORS[platform]
                        badge_parts.append(
                            f'<span style="display:inline-flex;align-items:center;gap:5px;'
                            f'background:{c}15;color:{c};border-radius:8px;'
                            f'padding:5px 10px;font-size:12px;font-weight:600;letter-spacing:.01em;">'
                            f'{PLATFORM_ICON_HTML[platform]}{label}</span>'
                        )
                badges_html = " ".join(badge_parts)

                favicon = (
                    f'<img src="{img_url}" class="ts-tile-fav" '
                    f'onerror="this.style.display=\'none\'">'
                    if img_url else
                    '<div class="ts-tile-fav" style="display:flex;align-items:center;'
                    'justify-content:center;font-size:16px;">🏢</div>'
                )

                with col, st.container(key=f"tile-{row_start}-{col_idx}"):
                    st.markdown(f"""
                    <div class="ts-tile">
                      <div class="ts-tile-top">
                        {favicon}
                        <div class="ts-tile-name">{client['bedrijfsnaam']}</div>
                      </div>
                      <div class="ts-tile-badges">{badges_html}</div>
                      <div class="ts-tile-foot">{total} posts / week</div>
                    </div>
                    """, unsafe_allow_html=True)

                    with st.expander("Details"):
                        col_a, col_b = st.columns(2)
                        with col_a:
                            st.markdown("**Platformen**")
                            for platform, label in PLATFORM_LABELS.items():
                                count = client.get(f"{platform}_posts_pw", 0)
                                if count:
                                    foll = followers.get(platform, "—")
                                    foll_text = f"· {foll}" if foll != "—" else ""
                                    st.markdown(
                                        f'{PLATFORM_ICON_HTML[platform]}'
                                        f'<span style="color:{PLATFORM_COLORS[platform]};font-weight:600;">'
                                        f'{label}</span> <span style="color:#555;">'
                                        f'{count}x/week {foll_text}</span>',
                                        unsafe_allow_html=True,
                                    )
                            st.markdown("**Toon**")
                            st.caption(client.get("toon") or "_—_")
                            st.markdown("**Doelgroep**")
                            st.caption(client.get("doelgroep") or "_—_")
                        with col_b:
                            st.markdown("**Kernthema's**")
                            themas = client.get("kernthemas", "")
                            if themas:
                                for t in themas.split(","):
                                    st.caption(f"· {t.strip()}")
                            else:
                                st.caption("_—_")
                            st.markdown("**Hashtags**")
                            st.code(client.get("vaste_hashtags") or "—", language=None)
                            urls = {
                                "Website":   client.get("website_url"),
                                "Instagram": client.get("instagram_url"),
                                "LinkedIn":  client.get("linkedin_url"),
                                "Facebook":  client.get("facebook_url"),
                            }
                            links = [(lbl, url) for lbl, url in urls.items() if url]
                            if links:
                                st.markdown("**Links**")
                                for lbl, url in links:
                                    st.markdown(f"[{lbl}]({url})")

    st.caption(f"Gegevens worden elke 5 minuten vernieuwd · Laatste update: {_now_ams().strftime('%H:%M:%S')}")

# ── Goedkeuring tab ───────────────────────────────────────────────────────────

STATUS_OPTIONS  = ["concept", "goedgekeurd", "afgewezen"]
STATUS_COLORS   = {"concept": "#f59e0b", "goedgekeurd": "#22c55e", "afgewezen": "#ef4444"}
STATUS_LABELS   = {"concept": "⏳ Concept", "goedgekeurd": "✅ Goedgekeurd", "afgewezen": "❌ Afgewezen"}

# ── Planning-tab: kolommen K-Q in Posts_YYYY_WNN ─────────────────────────────
PLANNING_HEADERS = [
    "geplande_datum", "geplande_tijd", "afbeelding_url", "afbeelding_drive_id",
    "publicatie_status", "meta_post_id", "publicatie_log",
]
PLANNING_COL_LETTERS = {
    "geplande_datum": "K", "geplande_tijd": "L", "afbeelding_url": "M",
    "afbeelding_drive_id": "N", "publicatie_status": "O",
    "meta_post_id": "P", "publicatie_log": "Q",
}
PUB_STATUS_COLORS = {
    "": "#9ca3af", "gepland": "#3b82f6", "bezig": "#f59e0b",
    "gepubliceerd": "#22c55e", "mislukt": "#ef4444",
}
PUB_STATUS_LABELS = {
    "": "— Niet ingepland", "gepland": "🕒 Gepland", "bezig": "⏳ Bezig met publiceren",
    "gepubliceerd": "✅ Gepubliceerd", "mislukt": "❌ Mislukt",
}


def _get_write_client(sa_info: dict):
    creds = WriteCredentials.from_service_account_info(sa_info, scopes=WRITE_SCOPES)
    return gspread.authorize(creds)


def _get_read_client(sa_info: dict):
    creds = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
    return gspread.authorize(creds)


@st.cache_data(ttl=60)
def load_post_tabs(spreadsheet_id: str, sa_info_json: str) -> list[str]:
    """Geeft alle tabbladen terug die beginnen met 'Posts_'."""
    sa_info = json.loads(sa_info_json)
    gc = _get_read_client(sa_info)
    spreadsheet = gc.open_by_key(spreadsheet_id)
    return sorted(
        [ws.title for ws in spreadsheet.worksheets() if ws.title.startswith("Posts_")],
        reverse=True,
    )


@st.cache_data(ttl=30)
def load_posts_from_tab(spreadsheet_id: str, tab_name: str, sa_info_json: str) -> list[dict]:
    """Laadt alle posts uit een specifiek tabblad."""
    sa_info = json.loads(sa_info_json)
    gc = _get_read_client(sa_info)
    worksheet = gc.open_by_key(spreadsheet_id).worksheet(tab_name)
    return worksheet.get_all_records(default_blank="")


def save_titles(spreadsheet_id: str, tab_name: str, titles: dict, sa_info_json: str):
    """Schrijft beeldtitels naar kolom J. titles = {row_index: titel_str}"""
    if not titles:
        return
    sa_info = json.loads(sa_info_json)
    gc = _get_write_client(sa_info)
    worksheet = gc.open_by_key(spreadsheet_id).worksheet(tab_name)
    batch = []
    headers = worksheet.row_values(1)
    if len(headers) < 10 or headers[9] != "beeldtitel":
        batch.append({"range": "J1", "values": [["beeldtitel"]]})
    for ri, t in titles.items():
        batch.append({"range": f"J{ri}", "values": [[t]]})
    if batch:
        worksheet.batch_update(batch, value_input_option="RAW")


def generate_image_titles(rows: list, api_key: str) -> dict:
    """Genereert beeldtitels (max 6 woorden) via Claude Haiku. Geeft {row_idx: titel} terug."""
    import anthropic
    ac = anthropic.Anthropic(api_key=api_key)
    results = {}
    for row_idx, post in rows:
        prompt = (
            f"Schrijf een beeldtitel voor deze social media post. "
            f"Maximaal 6 woorden. Alleen de titel, geen uitleg, geen aanhalingstekens.\n\n"
            f"Platform: {post.get('platform', '')}\n"
            f"Tekst: {post.get('caption', '')}"
        )
        msg = ac.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=30,
            messages=[{"role": "user", "content": prompt}],
        )
        title = msg.content[0].text.strip().strip('"').strip("'")
        words = title.split()
        results[row_idx] = " ".join(words[:6])
    return results


def save_statuses(spreadsheet_id: str, tab_name: str, updates: dict, sa_info_json: str):
    """Schrijft statuswijzigingen terug naar het tabblad. updates = {row_index: (status, opmerking)}"""
    sa_info = json.loads(sa_info_json)
    gc = _get_write_client(sa_info)
    worksheet = gc.open_by_key(spreadsheet_id).worksheet(tab_name)
    batch = []
    for row_idx, (status, opmerking) in updates.items():
        batch.append({"range": f"H{row_idx}:I{row_idx}", "values": [[status, opmerking]]})
    worksheet.batch_update(batch, value_input_option="RAW")


def _ensure_planning_columns(worksheet):
    """Vult de header-rij aan met de planningskolommen (K-Q) als die nog ontbreken.
    Oudere Posts_*-tabbladen (geüpload vóór deze feature) hebben alleen kolom A-J."""
    headers = worksheet.row_values(1)
    if len(headers) >= 17 and headers[10:17] == PLANNING_HEADERS:
        return
    headers = headers[:10]
    while len(headers) < 10:
        headers.append("")
    headers = headers + PLANNING_HEADERS
    worksheet.update(range_name="A1", values=[headers], value_input_option="RAW")
    worksheet.format("A1:Q1", {"textFormat": {"bold": True}})


def save_planning_fields(spreadsheet_id: str, tab_name: str, row_idx: int, fields: dict, sa_info_json: str):
    """Schrijft één of meer planningsvelden voor een rij weg.
    fields = {"geplande_datum": ..., "publicatie_status": ..., ...} (keys uit PLANNING_HEADERS)."""
    if not fields:
        return
    sa_info = json.loads(sa_info_json)
    gc = _get_write_client(sa_info)
    worksheet = gc.open_by_key(spreadsheet_id).worksheet(tab_name)
    _ensure_planning_columns(worksheet)
    batch = []
    for key, val in fields.items():
        col = PLANNING_COL_LETTERS[key]
        batch.append({"range": f"{col}{row_idx}", "values": [[val]]})
    worksheet.batch_update(batch, value_input_option="RAW")


MONTHS_NL_LONG = [
    "januari", "februari", "maart", "april", "mei", "juni",
    "juli", "augustus", "september", "oktober", "november", "december",
]

PLATFORM_LABELS_EXPORT = {"instagram": "Instagram", "linkedin": "LinkedIn", "facebook": "Facebook"}
PLATFORM_COLORS_EXPORT  = {"instagram": (0xE1, 0x30, 0x6C), "linkedin": (0x00, 0x77, 0xB5), "facebook": (0x18, 0x77, 0xF2)}


@st.cache_data(ttl=300)
def load_client_dict(spreadsheet_id: str, sa_info_json: str) -> dict:
    """Laadt klantprofielen als dict keyed by bedrijfsnaam."""
    clients = load_clients()
    return {c["bedrijfsnaam"]: c for c in clients}


@st.cache_data(ttl=300)
def load_meta_publish_context(spreadsheet_id: str, sa_info_json: str):
    """Laadt {klant_id: {ig_id, page_id}} en Facebook Page Access Tokens, nodig voor
    de "🚀 Nu publiceren"-knop in de Planning-tab. Geeft (accounts, page_tokens) terug;
    bij ontbrekend META_ACCESS_TOKEN of een fout: ({}, {})."""
    token = os.getenv("META_ACCESS_TOKEN") or st.secrets.get("META_ACCESS_TOKEN")
    if not token:
        return {}, {}
    try:
        sa_info = json.loads(sa_info_json)
        creds = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        gc = gspread.authorize(creds)
        spreadsheet = gc.open_by_key(spreadsheet_id)
        accounts = pub._load_meta_accounts(spreadsheet)
        page_tokens = pub._page_access_tokens(token)
        return accounts, page_tokens
    except Exception:
        return {}, {}


def _regenerate_prompt(post: dict, client: dict) -> str:
    opmerking = post.get("opmerkingen", "").strip()
    return f"""De volgende social media post voor {client.get('bedrijfsnaam','')} werd afgewezen.

Originele post:
Platform: {post.get('platform','')}
Dag: {post.get('dag','')} ({post.get('publicatiedatum','')})
Tekst: {post.get('caption','')}
Hashtags: {post.get('hashtags','')}

Reden van afwijzing: {opmerking or 'Geen reden opgegeven — verbeter de kwaliteit en relevantie.'}

Klantprofiel:
- Toon: {client.get('toon','')}
- Doelgroep: {client.get('doelgroep','')}
- Kernthema's: {client.get('kernthemas','')}
- Vaste hashtags: {client.get('vaste_hashtags','')}
- Vermijd: {client.get('vermijd','')}

Schrijf een nieuwe, verbeterde versie die de feedback adresseert.
Regels: geen koppeltekens (-) in de tekst. CTA linkt naar een relevante pagina op {client.get('website_url','de website')}.

Geef alleen valide JSON terug:
{{"caption": "...", "hashtags": "..."}}"""


def regenerate_rejected(all_posts: list[dict], client_dict: dict, api_key: str,
                        spreadsheet_id: str, tab_name: str, sa_info_json: str,
                        filter_bedrijfsnaam: str | None = None) -> tuple[int, str]:
    """
    Regenereert afgewezen posts via Claude API en schrijft ze terug naar de sheet.
    Geeft (aantal_verwerkt, foutmelding) terug.
    Rij-indices worden berekend op de volledige lijst zodat sheet-nummers kloppen.
    """
    import anthropic

    rejected = [
        (i + 2, p) for i, p in enumerate(all_posts)
        if p.get("status") == "afgewezen"
        and (filter_bedrijfsnaam is None or p.get("bedrijfsnaam") == filter_bedrijfsnaam)
    ]
    if not rejected:
        return 0, ""

    try:
        ac = anthropic.Anthropic(api_key=api_key)
        sa_info = json.loads(sa_info_json)
        gc = _get_write_client(sa_info)
        worksheet = gc.open_by_key(spreadsheet_id).worksheet(tab_name)

        updates = []
        for row_idx, post in rejected:
            client = client_dict.get(post.get("bedrijfsnaam", ""), {})
            prompt = _regenerate_prompt(post, client)

            msg = ac.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                system="Je bent een professionele social media contentschrijver. Retourneer uitsluitend valide JSON.",
                messages=[{"role": "user", "content": prompt}],
            )
            raw = msg.content[0].text.strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1].lstrip("json").strip()
            raw = raw.rstrip("```").strip()
            new_post = json.loads(raw)

            updates.append({
                "range": f"F{row_idx}:H{row_idx}",
                "values": [[new_post.get("caption", ""), new_post.get("hashtags", ""), "concept"]],
            })

        if updates:
            worksheet.batch_update(updates, value_input_option="RAW")

        return len(rejected), ""

    except Exception as e:
        return 0, str(e)


def regenerate_single_post(post: dict, client_dict: dict, api_key: str,
                           spreadsheet_id: str, tab_name: str, sa_info_json: str,
                           row_idx: int) -> tuple[dict | None, str]:
    """
    Genereert direct een nieuw concept voor één afgewezen post (op basis van de
    opgegeven reden) en schrijft het meteen terug naar de sheet (status -> concept).
    Geeft (nieuwe_post_velden | None, foutmelding) terug.
    """
    import anthropic

    try:
        ac = anthropic.Anthropic(api_key=api_key)
        client = client_dict.get(post.get("bedrijfsnaam", ""), {})
        prompt = _regenerate_prompt(post, client)

        msg = ac.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            system="Je bent een professionele social media contentschrijver. Retourneer uitsluitend valide JSON.",
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1].lstrip("json").strip()
        raw = raw.rstrip("```").strip()
        new_post = json.loads(raw)

        new_caption  = new_post.get("caption", "")
        new_hashtags = new_post.get("hashtags", "")

        sa_info = json.loads(sa_info_json)
        gc = _get_write_client(sa_info)
        worksheet = gc.open_by_key(spreadsheet_id).worksheet(tab_name)
        worksheet.batch_update([{
            "range": f"F{row_idx}:H{row_idx}",
            "values": [[new_caption, new_hashtags, "concept"]],
        }], value_input_option="RAW")

        return {"caption": new_caption, "hashtags": new_hashtags}, ""
    except Exception as e:
        return None, str(e)


def _build_approved_docx(bedrijfsnaam: str, approved_posts: list[dict]) -> bytes:
    """Bouwt een Word-document met alleen de goedgekeurde posts."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin   = Pt(70)
        section.right_margin  = Pt(70)

    title = doc.add_heading(f"{bedrijfsnaam} — Definitieve Social Media Posts", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph(f"Gegenereerd op {_now_ams().strftime('%d-%m-%Y %H:%M')} · Alleen goedgekeurde posts")
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for platform in ("instagram", "linkedin", "facebook"):
        platform_posts = [p for p in approved_posts if p.get("platform") == platform]
        if not platform_posts:
            continue

        label = PLATFORM_LABELS_EXPORT[platform]
        r, g, b = PLATFORM_COLORS_EXPORT[platform]
        color = RGBColor(r, g, b)

        heading = doc.add_heading(f"{label} — {len(platform_posts)} post{'s' if len(platform_posts) > 1 else ''}", level=2)
        for run in heading.runs:
            run.font.color.rgb = color

        for post in platform_posts:
            dag_p = doc.add_paragraph()
            dag_run = dag_p.add_run(f"📅 {post.get('dag','')} — {post.get('publicatiedatum','')}")
            dag_run.bold = True
            dag_run.font.size = Pt(11)

            beeldtitel = post.get("beeldtitel", "")
            if beeldtitel:
                bt_p = doc.add_paragraph()
                bt_p.add_run("🖼️ Beeldtitel: ").bold = True
                bt_run = bt_p.add_run(beeldtitel)
                bt_run.font.color.rgb = color
                bt_run.font.size = Pt(11)

            doc.add_paragraph(post.get("caption", ""))

            ht_p = doc.add_paragraph()
            ht_run = ht_p.add_run(post.get("hashtags", ""))
            ht_run.font.color.rgb = color
            ht_run.font.size = Pt(10)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _eff_status(cur_updates, row_idx, post):
    ov = cur_updates.get(row_idx)
    return ov[0] if ov else post.get("status", "concept")


def _eff_note(cur_updates, row_idx, post):
    ov = cur_updates.get(row_idx)
    return ov[1] if ov else post.get("opmerkingen", "")


def _render_post_card(selected_tab, row_idx, post, color, client_dict, spreadsheet_id, sa_json):
    """Rendert één post-kaart. Maakt deel uit van de render_approval_interface-fragment,
    zodat een klik op ✅/❌ ook de voortgangsbalk, metrics en actieknoppen direct bijwerkt."""
    state_key    = f"updates_{selected_tab}"
    pending_key  = f"pending_{selected_tab}"
    titles_key   = f"titles_{selected_tab}"
    ptitles_key  = f"ptitles_{selected_tab}"
    regen_key    = f"regen_{selected_tab}"
    cur_updates    = st.session_state[state_key]
    pending        = st.session_state[pending_key]
    cur_titles     = st.session_state[titles_key]
    pending_titles = st.session_state[ptitles_key]
    if regen_key not in st.session_state:
        st.session_state[regen_key] = {}
    regen_overrides = st.session_state[regen_key]

    cur_status = _eff_status(cur_updates, row_idx, post)

    # Toon eventueel net geregenereerde tekst i.p.v. de originele post-inhoud
    overridden  = regen_overrides.get(row_idx)
    disp_caption  = overridden["caption"]  if overridden else post.get("caption", "")
    disp_hashtags = overridden["hashtags"] if overridden else post.get("hashtags", "")

    col_post, col_ctrl = st.columns([5, 3])

    with col_post:
        if overridden:
            st.caption("🔄 Nieuw concept gegenereerd op basis van je feedback")
        st.markdown(
            f'<p style="font-weight:600;font-size:13px;margin:4px 0 2px 0;">'
            f'📅 {post.get("dag","")} — {post.get("publicatiedatum","")}</p>'
            f'<div style="background:#f8f8f8;border-left:3px solid {color};'
            f'padding:8px 12px;border-radius:0 8px 8px 0;font-size:13px;white-space:pre-wrap;">'
            f'{disp_caption}</div>'
            f'<p style="font-size:11px;color:{color};margin:3px 0 4px 0;">'
            f'{disp_hashtags}</p>',
            unsafe_allow_html=True,
        )
        current_title = cur_titles.get(row_idx) or post.get("beeldtitel", "")
        new_title = st.text_input(
            "Beeldtitel",
            value=current_title,
            key=f"title_{row_idx}",
            placeholder="Max 6 woorden voor de afbeelding...",
            label_visibility="collapsed",
        )
        word_count = len(new_title.split()) if new_title.strip() else 0
        if word_count > 6:
            st.caption(f"⚠️ {word_count}/6 woorden — wordt ingekort bij opslaan")
        elif new_title.strip():
            st.caption(f"🖼️ {word_count}/6 woorden")
        if new_title != current_title:
            cur_titles[row_idx]    = new_title
            pending_titles[row_idx] = " ".join(new_title.split()[:6])

    with col_ctrl:
        # Gekleurde status-badge — in een placeholder zodat we 'm na een klik
        # direct kunnen vervangen zonder de hele (sub)pagina opnieuw te laten
        # draaien (dat gaf juist de "hapering" terug die we al hadden opgelost).
        def _badge_html(status):
            b = {
                "goedgekeurd": ("#22c55e", "✅ Goedgekeurd"),
                "afgewezen":   ("#ef4444", "❌ Afgewezen"),
                "concept":     ("#f59e0b", "⏳ Concept"),
            }.get(status, ("#999", status))
            return (
                f'<div style="background:{b[0]};color:#fff;font-weight:700;'
                f'text-align:center;border-radius:8px;padding:5px 8px;'
                f'font-size:12px;margin-bottom:8px;">{b[1]}</div>'
            )

        badge_slot = st.empty()
        badge_slot.markdown(_badge_html(cur_status), unsafe_allow_html=True)

        col_ok, col_rej_btn = st.columns(2)
        with col_ok:
            if st.button("✅", key=f"ok_{row_idx}", use_container_width=True,
                         disabled=cur_status == "goedgekeurd"):
                cur_updates[row_idx] = ("goedgekeurd", "")
                pending[row_idx]     = ("goedgekeurd", "")
                st.rerun(scope="fragment")
        with col_rej_btn:
            if st.button("❌", key=f"rej_{row_idx}", use_container_width=True,
                         disabled=cur_status == "afgewezen"):
                cur_updates[row_idx] = ("afgewezen", _eff_note(cur_updates, row_idx, post))
                pending[row_idx]     = ("afgewezen", _eff_note(cur_updates, row_idx, post))
                st.rerun(scope="fragment")

        if cur_status == "afgewezen":
            typed = st.text_input(
                "Reden",
                value=_eff_note(cur_updates, row_idx, post),
                key=f"note_{row_idx}",
                placeholder="Wat moet er anders?",
                label_visibility="collapsed",
            )
            cur_updates[row_idx] = ("afgewezen", typed)
            pending[row_idx]     = ("afgewezen", typed)

            if st.button(
                "🔄 Genereer nieuw concept",
                key=f"autoregen_{row_idx}",
                use_container_width=True,
                disabled=not typed.strip(),
                help="Schrijft direct een nieuwe versie op basis van je feedback en zet de status terug naar concept"
                     if typed.strip() else "Vul eerst een reden in zodat de AI weet wat er anders moet",
            ):
                api_key = os.getenv("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY", "")
                if not api_key:
                    st.error("ANTHROPIC_API_KEY ontbreekt.")
                else:
                    with st.spinner("Nieuw concept genereren..."):
                        post_for_prompt = {**post, "caption": disp_caption,
                                           "hashtags": disp_hashtags, "opmerkingen": typed}
                        new_fields, err = regenerate_single_post(
                            post_for_prompt, client_dict, api_key,
                            spreadsheet_id, selected_tab, sa_json, row_idx,
                        )
                    if err:
                        st.error(f"Fout bij regenereren: {err}")
                    else:
                        regen_overrides[row_idx] = new_fields
                        cur_updates[row_idx] = ("concept", "")
                        pending.pop(row_idx, None)
                        load_posts_from_tab.clear()
                        st.success("✓ Nieuw concept klaar — beoordeel hierboven")
                        st.rerun(scope="fragment")

    st.markdown(
        "<hr style='margin:4px 0 10px 0;border:none;border-top:1px solid #eee;'>",
        unsafe_allow_html=True,
    )


def _render_planning_card(row_idx, post, bedrijfsnaam, client_dict, spreadsheet_id, selected_tab, sa_json,
                           accounts=None, page_tokens=None):
    """Eén kaart per goedgekeurde post: afbeelding uploaden + datum/tijd inplannen."""
    platform = post.get("platform", "")
    color = PLATFORM_COLORS.get(platform, "#999")
    icon = PLATFORM_ICON_HTML.get(platform, "")

    col_post, col_img, col_plan = st.columns([4, 2, 3])

    with col_post:
        st.markdown(
            f'<p style="font-weight:600;font-size:13px;margin:4px 0 2px 0;">'
            f'{icon}📅 {post.get("dag","")} — {post.get("publicatiedatum","")}</p>'
            f'<div style="background:#f8f8f8;border-left:3px solid {color};'
            f'padding:8px 12px;border-radius:0 8px 8px 0;font-size:13px;white-space:pre-wrap;'
            f'max-height:120px;overflow-y:auto;">'
            f'{post.get("caption","")}</div>'
            f'<p style="font-size:11px;color:{color};margin:3px 0 4px 0;">{post.get("hashtags","")}</p>',
            unsafe_allow_html=True,
        )
        if post.get("beeldtitel"):
            st.caption(f"🖼️ {post['beeldtitel']}")

    afbeelding_url = post.get("afbeelding_url", "")
    with col_img:
        if afbeelding_url:
            st.image(afbeelding_url, use_container_width=True)
        uploaded = st.file_uploader(
            "Afbeelding",
            type=["png", "jpg", "jpeg"],
            key=f"upload_{selected_tab}_{row_idx}",
            label_visibility="collapsed",
        )
        if uploaded is not None:
            folder_id = client_dict.get(bedrijfsnaam, {}).get("google_doc_folder_id", "").strip()
            if not folder_id:
                st.error("Geen Drive-map gekoppeld aan deze klant (kolom 'google_doc_folder_id' ontbreekt).")
            else:
                with st.spinner("Uploaden naar Drive..."):
                    try:
                        sa_info = json.loads(sa_json)
                        result = drive_upload.upload_image(
                            sa_info, folder_id, uploaded.name,
                            uploaded.getvalue(), uploaded.type or "image/jpeg",
                        )
                        save_planning_fields(
                            spreadsheet_id, selected_tab, row_idx,
                            {"afbeelding_url": result["url"], "afbeelding_drive_id": result["id"]},
                            sa_json,
                        )
                        load_posts_from_tab.clear()
                        st.success("✓ Afbeelding geüpload")
                        st.rerun(scope="fragment")
                    except Exception as e:
                        st.error(f"Upload mislukt: {e}")

    pub_status = post.get("publicatie_status", "")
    with col_plan:
        if pub_status in ("gepland", "bezig", "gepubliceerd"):
            badge_color = PUB_STATUS_COLORS.get(pub_status, "#9ca3af")
            badge_label = PUB_STATUS_LABELS.get(pub_status, pub_status)
            st.markdown(
                f'<div style="background:{badge_color};color:#fff;font-weight:700;'
                f'text-align:center;border-radius:8px;padding:5px 8px;font-size:12px;'
                f'margin-bottom:8px;">{badge_label}</div>',
                unsafe_allow_html=True,
            )
            if post.get("geplande_datum"):
                st.caption(f"🕒 {post['geplande_datum']} om {post.get('geplande_tijd','')}")
            if pub_status == "gepubliceerd" and post.get("meta_post_id"):
                st.caption(f"Post-ID: {post['meta_post_id']}")
            if pub_status == "gepland":
                if st.button("❌ Annuleer planning", key=f"cancel_{selected_tab}_{row_idx}",
                              use_container_width=True):
                    save_planning_fields(
                        spreadsheet_id, selected_tab, row_idx,
                        {"geplande_datum": "", "geplande_tijd": "", "publicatie_status": ""},
                        sa_json,
                    )
                    load_posts_from_tab.clear()
                    st.rerun(scope="fragment")
        else:
            if pub_status == "mislukt":
                st.markdown(
                    f'<div style="background:{PUB_STATUS_COLORS["mislukt"]};color:#fff;font-weight:700;'
                    f'text-align:center;border-radius:8px;padding:5px 8px;font-size:12px;'
                    f'margin-bottom:6px;">{PUB_STATUS_LABELS["mislukt"]}</div>',
                    unsafe_allow_html=True,
                )
                if post.get("publicatie_log"):
                    st.caption(f"⚠️ {post['publicatie_log']}")

            default_date_str = post.get("geplande_datum") or ""
            try:
                default_date = date.fromisoformat(default_date_str) if default_date_str else date.today()
            except ValueError:
                default_date = date.today()

            default_time_str = post.get("geplande_tijd") or "09:00"
            try:
                h, m = default_time_str.split(":")
                default_time = dtime(int(h), int(m))
            except (ValueError, TypeError):
                default_time = dtime(9, 0)

            d = st.date_input("Datum", value=default_date, key=f"date_{selected_tab}_{row_idx}",
                               label_visibility="collapsed")
            t = st.time_input("Tijd", value=default_time, key=f"time_{selected_tab}_{row_idx}",
                               label_visibility="collapsed")

            if not afbeelding_url:
                st.caption("⚠️ Upload eerst een afbeelding")

            if st.button("📌 Inplannen", key=f"plan_{selected_tab}_{row_idx}",
                          use_container_width=True, disabled=not afbeelding_url):
                save_planning_fields(
                    spreadsheet_id, selected_tab, row_idx,
                    {
                        "geplande_datum": d.isoformat(),
                        "geplande_tijd": t.strftime("%H:%M"),
                        "publicatie_status": "gepland",
                        "publicatie_log": "",
                    },
                    sa_json,
                )
                load_posts_from_tab.clear()
                st.success("✓ Ingepland")
                st.rerun(scope="fragment")

        if platform in ("instagram", "facebook") and afbeelding_url and pub_status not in ("gepubliceerd", "bezig"):
            meta_token = os.getenv("META_ACCESS_TOKEN") or st.secrets.get("META_ACCESS_TOKEN")
            klant_id = post.get("klant_id", "")
            account = (accounts or {}).get(klant_id, {})
            if st.button("🚀 Nu publiceren", key=f"publish_now_{selected_tab}_{row_idx}",
                          use_container_width=True):
                if not meta_token:
                    st.error("META_ACCESS_TOKEN ontbreekt — kan niet publiceren.")
                else:
                    save_planning_fields(
                        spreadsheet_id, selected_tab, row_idx,
                        {"publicatie_status": "bezig"}, sa_json,
                    )
                    try:
                        post_id = pub.publish_post(post, account, page_tokens or {}, meta_token)
                    except RuntimeError as e:
                        save_planning_fields(
                            spreadsheet_id, selected_tab, row_idx,
                            {"publicatie_status": "mislukt", "publicatie_log": str(e)},
                            sa_json,
                        )
                        load_posts_from_tab.clear()
                        st.error(f"Publiceren mislukt: {e}")
                    else:
                        now = _now_ams()
                        save_planning_fields(
                            spreadsheet_id, selected_tab, row_idx,
                            {
                                "publicatie_status": "gepubliceerd",
                                "meta_post_id": post_id,
                                "publicatie_log": "",
                                "geplande_datum": now.date().isoformat(),
                                "geplande_tijd": now.strftime("%H:%M"),
                            },
                            sa_json,
                        )
                        load_posts_from_tab.clear()
                        st.success("✓ Direct gepubliceerd")
                        st.rerun(scope="fragment")

    st.markdown(
        "<hr style='margin:4px 0 10px 0;border:none;border-top:1px solid #eee;'>",
        unsafe_allow_html=True,
    )


@st.fragment
def render_planning_interface(posts, client_dict, spreadsheet_id, selected_tab, sa_json):
    """Plan-tab: afbeeldingen uploaden, datum/tijd inplannen voor goedgekeurde posts, of
    direct publiceren via "🚀 Nu publiceren". Geplande posts worden verder afgehandeld door
    systems/publish_scheduled_posts.py (GitHub Actions cron)."""
    accounts, page_tokens = load_meta_publish_context(spreadsheet_id, sa_json)
    planbaar = [(i, p) for i, p in enumerate(posts, start=2) if p.get("status") == "goedgekeurd"]

    if not planbaar:
        st.info(
            "Nog geen goedgekeurde posts in deze week. Keur eerst posts goed in de tab "
            "'✅ Goedkeuring' — daarna verschijnen ze hier om in te plannen."
        )
        return

    # ── Agenda-overzicht: alle ingeplande posts, gesorteerd op datum/tijd ───────
    geplande = [(i, p) for i, p in planbaar if p.get("geplande_datum")]
    if geplande:
        st.markdown('<p class="ts-stat-title">🗓️ Agenda — ingeplande posts</p>', unsafe_allow_html=True)
        for i, p in sorted(geplande, key=lambda ip: (ip[1].get("geplande_datum", ""), ip[1].get("geplande_tijd", ""))):
            platform = p.get("platform", "")
            icon = PLATFORM_ICON_HTML.get(platform, "")
            pub_status = p.get("publicatie_status", "")
            badge_color = PUB_STATUS_COLORS.get(pub_status, "#9ca3af")
            badge_label = PUB_STATUS_LABELS.get(pub_status, pub_status)
            afbeelding_url = p.get("afbeelding_url", "")
            thumb_html = (
                f'<img class="ts-post-thumb" src="{afbeelding_url}" alt="">'
                if afbeelding_url else '<div class="ts-post-thumb ts-post-thumb-empty">🖼️</div>'
            )
            st.markdown(
                f'<div class="ts-post-card" style="display:flex;align-items:center;gap:12px;">'
                f'{thumb_html}'
                f'<div style="flex:1;">'
                f'{icon}<span style="font-weight:600;">{p.get("bedrijfsnaam","")}</span>'
                f'<div style="font-size:12px;color:#666;">{p.get("geplande_datum","")} · {p.get("geplande_tijd","")}</div>'
                f'</div>'
                f'<div style="background:{badge_color};color:#fff;font-weight:700;border-radius:8px;'
                f'padding:4px 10px;font-size:12px;white-space:nowrap;">{badge_label}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if pub_status == "mislukt" and p.get("publicatie_log"):
                st.caption(f"⚠️ {p['publicatie_log']}")
        st.divider()

    # ── Per klant: posts uploaden/inplannen ──────────────────────────────────
    st.markdown('<p class="ts-stat-title">📤 Afbeeldingen uploaden & inplannen</p>', unsafe_allow_html=True)

    by_client: dict[str, list] = {}
    for i, p in planbaar:
        by_client.setdefault(p.get("bedrijfsnaam", "Onbekend"), []).append((i, p))

    for naam, rows in by_client.items():
        with st.expander(f"{naam} ({len(rows)} post{'s' if len(rows) != 1 else ''})", expanded=(len(by_client) == 1)):
            for i, p in rows:
                _render_planning_card(i, p, naam, client_dict, spreadsheet_id, selected_tab, sa_json,
                                       accounts, page_tokens)


@st.fragment
def render_approval_interface(posts, client_dict, spreadsheet_id, selected_tab, sa_json):
    # Data is al geladen buiten de fragment — geen API-calls bij klikken

    if not posts:
        st.warning("Geen posts gevonden in dit tabblad.")
        return

    # ── Groepeer per klant ────────────────────────────────────────────────────
    clients_in_week: dict[str, list] = {}
    for i, post in enumerate(posts, start=2):
        name = post.get("bedrijfsnaam", "Onbekend")
        clients_in_week.setdefault(name, []).append((i, post))

    # ── Session state: statussen + beeldtitels ───────────────────────────────
    state_key    = f"updates_{selected_tab}"
    pending_key  = f"pending_{selected_tab}"
    titles_key   = f"titles_{selected_tab}"
    ptitles_key  = f"ptitles_{selected_tab}"
    if state_key   not in st.session_state: st.session_state[state_key]   = {}
    if pending_key not in st.session_state: st.session_state[pending_key] = {}
    if titles_key  not in st.session_state: st.session_state[titles_key]  = {}
    if ptitles_key not in st.session_state: st.session_state[ptitles_key] = {}
    cur_updates   = st.session_state[state_key]
    pending       = st.session_state[pending_key]
    cur_titles    = st.session_state[titles_key]
    pending_titles = st.session_state[ptitles_key]

    def _eff(row_idx, post):
        ov = cur_updates.get(row_idx)
        return ov[0] if ov else post.get("status", "concept")

    def _note_val(row_idx, post):
        ov = cur_updates.get(row_idx)
        return ov[1] if ov else post.get("opmerkingen", "")

    def _progress(rows):
        statuses = [_eff(ri, p) for ri, p in rows]
        approved = sum(1 for s in statuses if s == "goedgekeurd")
        if approved == len(statuses):
            return 100, "Klaar", "#22c55e"
        elif any(s != "concept" for s in statuses):
            return 66, "In review", "#f59e0b"
        return 0, "Niet gestart", "#ef4444"

    # ── Urgentie ──────────────────────────────────────────────────────────────
    is_urgent = _now_ams().weekday() in (3, 4)

    # ── Metrics ───────────────────────────────────────────────────────────────
    total_posts  = len(posts)
    approved_all = sum(_eff(i + 2, p) == "goedgekeurd" for i, p in enumerate(posts))
    rejected_all = sum(_eff(i + 2, p) == "afgewezen"   for i, p in enumerate(posts))
    pending_all  = total_posts - approved_all - rejected_all
    done_clients = sum(1 for rows in clients_in_week.values() if _progress(rows)[0] == 100)

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Klanten klaar", f"{done_clients}/{len(clients_in_week)}")
    c2.metric("Posts totaal", total_posts)
    c3.metric("✅ Goedgekeurd", approved_all)
    c4.metric("❌ Afgewezen", rejected_all)
    c5.metric("⏳ Concept", pending_all)

    overall_pct = int(approved_all / total_posts * 100) if total_posts else 0
    st.markdown(
        f'<div style="background:#e5e7eb;border-radius:99px;height:8px;margin:8px 0 16px 0;">'
        f'<div style="background:#22c55e;width:{overall_pct}%;height:8px;border-radius:99px;"></div></div>'
        f'<p style="font-size:12px;color:#666;margin-top:-8px;">{overall_pct}% van alle posts goedgekeurd</p>',
        unsafe_allow_html=True,
    )

    # ── Urgentie-banner ───────────────────────────────────────────────────────
    if is_urgent:
        incomplete = [n for n, rows in clients_in_week.items() if _progress(rows)[0] < 100]
        if incomplete:
            clr = "#ef4444" if _now_ams().weekday() == 4 else "#f59e0b"
            dag = "vrijdag" if _now_ams().weekday() == 4 else "donderdag"
            st.markdown(
                f'<div style="background:{clr}18;border:1.5px solid {clr};border-radius:10px;'
                f'padding:12px 16px;margin-bottom:12px;">'
                f'<span style="font-weight:700;color:{clr};">⚠️ Het is {dag} — '
                f'{len(incomplete)} klant{"en" if len(incomplete)>1 else ""} nog niet klaar.</span></div>',
                unsafe_allow_html=True,
            )

    # ── Statusoverzicht per klant ──────────────────────────────────────────────
    def _client_stats(rows):
        n_app = sum(1 for ri, p in rows if _eff(ri, p) == "goedgekeurd")
        n_rej = sum(1 for ri, p in rows if _eff(ri, p) == "afgewezen")
        n_con = len(rows) - n_app - n_rej
        return n_app, n_rej, n_con

    not_started, almost_done, waiting_regen = [], [], []
    for name, rows in clients_in_week.items():
        n_app, n_rej, n_con = _client_stats(rows)
        if n_app == 0 and n_rej == 0:
            not_started.append(name)
        elif n_rej > 0:
            waiting_regen.append((name, n_rej))
        elif n_app < len(rows):
            almost_done.append((name, n_app, len(rows)))

    with st.expander(
        f"📊 Statusoverzicht — {len(not_started)} nog niet gestart · "
        f"{len(waiting_regen)} met afgewezen posts · {len(almost_done)} bijna klaar",
        expanded=False,
    ):
        ov1, ov2, ov3 = st.columns(3)
        with ov1:
            st.markdown("**🔴 Nog niet gestart**")
            if not_started:
                for n in not_started:
                    st.markdown(f"- {n}")
            else:
                st.caption("Niemand — top!")
        with ov2:
            st.markdown("**🟠 Wacht op regeneratie**")
            if waiting_regen:
                for n, n_rej in waiting_regen:
                    st.markdown(f"- {n} _({n_rej} afgewezen)_")
            else:
                st.caption("Geen afgewezen posts open.")
        with ov3:
            st.markdown("**🟡 Bijna klaar**")
            if almost_done:
                for n, n_app, total in almost_done:
                    st.markdown(f"- {n} _({n_app}/{total} goedgekeurd)_")
            else:
                st.caption("Niets in dit overzicht.")

    st.divider()

    # ── Layout: navigator links, review rechts ────────────────────────────────
    # Volgorde eenmalig bepalen per tab zodat de lijst niet springt tijdens review
    sort_key = f"sort_{selected_tab}"
    if sort_key not in st.session_state:
        st.session_state[sort_key] = sorted(
            clients_in_week,
            key=lambda n: (_progress(clients_in_week[n])[0], n),
        )
    sorted_names = [n for n in st.session_state[sort_key] if n in clients_in_week]

    col_nav, col_review = st.columns([1, 3], gap="large")

    with col_nav:
        st.markdown("**Klanten**")
        nav_filter = st.selectbox(
            "Filter",
            ["Alle", "Nog niet gestart", "Heeft afgewezen posts", "In review", "Klaar"],
            label_visibility="collapsed",
            key=f"navfilter_{selected_tab}",
        )

        def _matches_filter(name):
            if nav_filter == "Alle":
                return True
            r = clients_in_week[name]
            n_app, n_rej, n_con = _client_stats(r)
            if nav_filter == "Nog niet gestart":
                return n_app == 0 and n_rej == 0
            if nav_filter == "Heeft afgewezen posts":
                return n_rej > 0
            if nav_filter == "In review":
                return 0 < n_app < len(r)
            if nav_filter == "Klaar":
                return n_app == len(r)
            return True

        filtered_names = [n for n in sorted_names if _matches_filter(n)]

        if not filtered_names:
            st.caption("Geen klanten in dit filter.")
            selected_client = sorted_names[0]
        else:
            radio_key = f"radio_client_{selected_tab}"
            current = st.session_state.get(radio_key)
            if current not in filtered_names:
                # Voorkom StreamlitAPIException wanneer de bewaarde selectie
                # niet meer in de gefilterde lijst voorkomt (filter gewijzigd)
                st.session_state.pop(radio_key, None)
            selected_client = st.radio(
                "Klant",
                filtered_names,
                format_func=lambda n: (
                    f"✅ {n}" if _progress(clients_in_week[n])[0] == 100
                    else f"{'🔴' if is_urgent else '⏳'} {n}"
                ),
                label_visibility="collapsed",
                key=radio_key,
            )

    with col_review:
        rows = clients_in_week[selected_client]

        # ── Auto-genereer beeldtitels voor posts zonder titel (eenmalig) ────
        api_key = os.getenv("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY", "")
        gen_flag = f"titles_generated_{selected_tab}_{selected_client}"
        if gen_flag not in st.session_state:
            st.session_state[gen_flag] = True
            if api_key:
                missing = [(ri, p) for ri, p in rows
                           if not cur_titles.get(ri) and not p.get("beeldtitel", "")]
                if missing:
                    with st.spinner("Beeldtitels genereren..."):
                        generated = generate_image_titles(missing, api_key)
                        save_titles(spreadsheet_id, selected_tab, generated, sa_json)
                        load_posts_from_tab.clear()
                    for ri, title in generated.items():
                        cur_titles[ri] = title

        pct, pct_label, pct_color = _progress(rows)
        n_app = sum(1 for ri, p in rows if _eff(ri, p) == "goedgekeurd")
        n_rej = sum(1 for ri, p in rows if _eff(ri, p) == "afgewezen")

        # Klant-header
        st.markdown(
            f'<div style="margin-bottom:14px;">'
            f'<div style="font-size:20px;font-weight:700;color:#18181b;">{selected_client}</div>'
            f'<div style="display:flex;align-items:center;gap:10px;margin-top:6px;">'
            f'<div style="flex:1;background:#e5e7eb;border-radius:99px;height:5px;">'
            f'<div style="background:{pct_color};width:{pct}%;height:5px;border-radius:99px;"></div></div>'
            f'<span style="font-size:12px;font-weight:700;color:{pct_color};">'
            f'{n_app}/{len(rows)} goedgekeurd · {pct_label}</span></div></div>',
            unsafe_allow_html=True,
        )

        # ── Posts per platform ────────────────────────────────────────────────
        for platform in ("instagram", "linkedin", "facebook"):
            platform_rows = [(ri, p) for ri, p in rows if p.get("platform") == platform]
            if not platform_rows:
                continue

            color = PLATFORM_COLORS[platform]
            icon  = PLATFORM_ICON_HTML[platform]
            st.markdown(
                f'<p style="font-weight:700;color:{color};margin:16px 0 8px 0;">'
                f'{icon}{PLATFORM_LABELS[platform]}</p>',
                unsafe_allow_html=True,
            )

            for row_idx, post in platform_rows:
                _render_post_card(selected_tab, row_idx, post, color, client_dict, spreadsheet_id, sa_json)

        # ── Actieknoppen onderaan de review ───────────────────────────────────
        client_row_indices = {ri for ri, _ in rows}
        client_pending        = {ri: v for ri, v in pending.items()        if ri in client_row_indices}
        client_pending_titles = {ri: v for ri, v in pending_titles.items() if ri in client_row_indices}
        has_pending = bool(client_pending or client_pending_titles)

        # ── Eénmalig de Word-export bouwen, hergebruikt door Download én Mail ──
        # (anders werd hetzelfde document twee keer opgebouwd bij elke klik)
        approved_posts = [
            {**p, "beeldtitel": cur_titles.get(ri) or p.get("beeldtitel", "")}
            for ri, p in rows if _eff(ri, p) == "goedgekeurd"
        ]
        approved_key = tuple(
            (ri, cur_titles.get(ri) or p.get("beeldtitel", ""))
            for ri, p in rows if _eff(ri, p) == "goedgekeurd"
        )
        docx_cache_key = f"docx_{selected_tab}_{selected_client}"
        cached = st.session_state.get(docx_cache_key)
        if approved_posts and (cached is None or cached[0] != approved_key):
            docx_bytes_shared = _build_approved_docx(selected_client, approved_posts)
            st.session_state[docx_cache_key] = (approved_key, docx_bytes_shared)
        elif approved_posts:
            docx_bytes_shared = cached[1]
        else:
            docx_bytes_shared = None
            st.session_state.pop(docx_cache_key, None)

        col_save, col_gen, col_regen, col_dl, col_mail = st.columns(5)

        with col_save:
            if st.button(
                f"💾 Opslaan ({len(client_pending) + len(client_pending_titles)})" if has_pending else "💾 Opgeslagen",
                key=f"save_{selected_client}",
                disabled=not has_pending,
                use_container_width=True,
                type="primary",
            ):
                if client_pending:
                    save_statuses(spreadsheet_id, selected_tab, client_pending, sa_json)
                    for ri in client_pending:
                        pending.pop(ri, None)
                if client_pending_titles:
                    save_titles(spreadsheet_id, selected_tab, client_pending_titles, sa_json)
                    for ri in client_pending_titles:
                        pending_titles.pop(ri, None)
                st.success("✓ Opgeslagen naar Google Sheets")

        with col_gen:
            if st.button(
                "🎨 Vernieuw titels",
                key=f"gen_titles_{selected_client}",
                use_container_width=True,
                help="Laat Claude nieuwe beeldtitels schrijven voor alle posts",
            ):
                api_key = os.getenv("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY", "")
                if not api_key:
                    st.error("ANTHROPIC_API_KEY ontbreekt.")
                else:
                    with st.spinner("Titels genereren..."):
                        generated = generate_image_titles(rows, api_key)
                        save_titles(spreadsheet_id, selected_tab, generated, sa_json)
                        load_posts_from_tab.clear()
                    for ri, title in generated.items():
                        cur_titles[ri] = title
                    st.success(f"✓ {len(generated)} titels vernieuwd en opgeslagen")

        with col_regen:
            if st.button(
                f"🔄 Regenereer ({n_rej})",
                key=f"regen_{selected_client}",
                disabled=n_rej == 0,
                use_container_width=True,
            ):
                api_key = os.getenv("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY", "")
                if not api_key:
                    st.error("ANTHROPIC_API_KEY ontbreekt.")
                else:
                    if client_pending:
                        save_statuses(spreadsheet_id, selected_tab, client_pending, sa_json)
                        for ri in client_pending:
                            pending.pop(ri, None)
                    load_posts_from_tab.clear()
                    fresh = load_posts_from_tab(spreadsheet_id, selected_tab, sa_json)
                    with st.spinner(f"{n_rej} posts regenereren..."):
                        count, err = regenerate_rejected(
                            fresh, client_dict, api_key,
                            spreadsheet_id, selected_tab, sa_json,
                            filter_bedrijfsnaam=selected_client,
                        )
                        load_posts_from_tab.clear()
                        # Ververs de posts in session state na regeneratie
                        st.session_state[f"posts_{selected_tab}"] = load_posts_from_tab(
                            spreadsheet_id, selected_tab, sa_json
                        )
                        st.session_state[state_key]   = {}
                        st.session_state[pending_key] = {}
                    if err:
                        st.error(f"Fout: {err}")
                    else:
                        st.success(f"✓ {count} posts herschreven")

        with col_dl:
            if not approved_posts:
                st.button("📄 Download", key=f"dl_{selected_client}", disabled=True, use_container_width=True)
            else:
                week_label = selected_tab.replace("Posts_", "").replace("_W", "_Week")
                st.download_button(
                    label=f"📄 Download ({len(approved_posts)})",
                    data=docx_bytes_shared,
                    file_name=f"{selected_client} - Definitief {week_label}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_btn_{selected_client}",
                    use_container_width=True,
                )

        with col_mail:
            from urllib.parse import quote as _quote
            week_label_mail = selected_tab.replace("Posts_", "").replace("_W", " week ")
            subject  = _quote(f"{selected_client} | Je kunt aan de slag met de afbeeldingen")
            body_txt = _quote(
                f"Hey,\n\n"
                f"We hebben de teksten voor {selected_client} voor {week_label_mail} goedgekeurd. "
                f"Kun je aan de slag met de afbeeldingen voor deze klant?\n\n"
                f"Alvast bedankt!"
            )
            mailto = f"mailto:studio@topmediagroep.nl?subject={subject}&body={body_txt}"

            # Let op: dit is een <button> in een eigen iframe (via components.html),
            # dus de globale .stButton-CSS (merklettertype, kleuren, randen) bereikt
            # 'm niet — we stylen 'm hier expliciet identiek aan st.button, en
            # geven het iframe géén eigen marge/border zodat 'ie netjes uitlijnt
            # met Download ernaast (zelfde hoogte, geen scrollbalk, géén canvas-rand).
            _btn_base_css = """
                width:100%; box-sizing:border-box; height:38.4px;
                padding:0 16px; font-size:14px; font-weight:600;
                border-radius:10px; font-family:'Plus Jakarta Sans',-apple-system,sans-serif;
                white-space:nowrap; line-height:1; display:inline-flex;
                align-items:center; justify-content:center; transition:all .12s ease;
            """
            _iframe_reset = """
                <style>
                    html, body { margin:0; padding:0; background:transparent; overflow:hidden; }
                </style>
            """
            if approved_posts:
                docx_b64 = base64.b64encode(docx_bytes_shared).decode()
                filename = f"{selected_client} - Definitief {week_label_mail}.docx"
                components.html(f"""
                {_iframe_reset}
                <button onclick="(function(){{
                    var a=document.createElement('a');
                    a.href='data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_b64}';
                    a.download='{filename}';
                    a.click();
                    setTimeout(function(){{window.location.href='{mailto}';}},600);
                }})()" style="{_btn_base_css}
                    background:#ffffff; color:#111827; border:1px solid #E5E7EB;
                    cursor:pointer;
                " onmouseover="this.style.borderColor='#4F46E5';this.style.color='#4F46E5';"
                  onmouseout="this.style.borderColor='#E5E7EB';this.style.color='#111827';"
                >📧&nbsp;Mail studio</button>
                """, height=39)
            else:
                components.html(f"""
                {_iframe_reset}
                <button disabled style="{_btn_base_css}
                    background:#F7F8FB; color:#9CA3AF; border:1px solid #E5E7EB;
                    cursor:not-allowed;
                ">📧&nbsp;Mail studio</button>
                """, height=39)


with tab_goedkeuring:
    st.subheader("Posts goedkeuren per week")

    spreadsheet_id = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID")
    sa_json = _sa_info_json()

    if not spreadsheet_id or not sa_json:
        st.error("Credentials ontbreken.")
    else:
        tabs = load_post_tabs(spreadsheet_id, sa_json)
        if not tabs:
            st.info("Nog geen posts beschikbaar. Voer eerst de pipeline uit.")
        else:
            selected_tab = st.selectbox(
                "Selecteer een week",
                tabs,
                format_func=lambda t: t.replace("Posts_", "").replace("_W", " · Week "),
            )
            # Laad posts — gebruik session state cache na regeneratie
            posts_ss_key = f"posts_{selected_tab}"
            if posts_ss_key in st.session_state:
                posts = st.session_state.pop(posts_ss_key)
            else:
                posts = load_posts_from_tab(spreadsheet_id, selected_tab, sa_json)

            # Filter posts op basis van rol
            if _logged_in_role != "admin":
                _assignments_g = load_medewerker_assignments(spreadsheet_id, sa_json)
                _allowed_g     = _assignments_g.get(_logged_in_user, [])
                if _allowed_g != "ALL":
                    posts = [p for p in posts if p.get("bedrijfsnaam", "") in _allowed_g]

            client_dict = load_client_dict(spreadsheet_id, sa_json)
            render_approval_interface(posts, client_dict, spreadsheet_id, selected_tab, sa_json)


# ── Planning tab ──────────────────────────────────────────────────────────────

with tab_planning:
    st.subheader("Content inplannen")
    st.caption(
        "Upload een afbeelding bij een goedgekeurde post en plan een datum/tijd in. "
        "Ingeplande posts worden automatisch gepubliceerd op Instagram en Facebook."
    )

    spreadsheet_id = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID")
    sa_json = _sa_info_json()

    if not spreadsheet_id or not sa_json:
        st.error("Credentials ontbreken.")
    else:
        tabs = load_post_tabs(spreadsheet_id, sa_json)
        if not tabs:
            st.info("Nog geen posts beschikbaar. Voer eerst de pipeline uit.")
        else:
            selected_tab = st.selectbox(
                "Selecteer een week",
                tabs,
                format_func=lambda t: t.replace("Posts_", "").replace("_W", " · Week "),
                key="planning_week_select",
            )
            posts = load_posts_from_tab(spreadsheet_id, selected_tab, sa_json)

            # Filter posts op basis van rol
            if _logged_in_role != "admin":
                _assignments_p = load_medewerker_assignments(spreadsheet_id, sa_json)
                _allowed_p     = _assignments_p.get(_logged_in_user, [])
                if _allowed_p != "ALL":
                    posts = [p for p in posts if p.get("bedrijfsnaam", "") in _allowed_p]

            client_dict = load_client_dict(spreadsheet_id, sa_json)
            render_planning_interface(posts, client_dict, spreadsheet_id, selected_tab, sa_json)


# ── Statistieken tab ──────────────────────────────────────────────────────────

with tab_statistieken:
    st.subheader("Realtime statistieken")
    st.caption(
        "Volgers, bereik, weergaven en engagement — opgehaald via de Meta Graph API "
        "(Instagram & Facebook). Klanten zonder koppeling tonen nog geen cijfers."
    )

    if not clients:
        st.warning("Geen klanten gevonden.")
    else:
        stats_history = load_statistics()
        latest_stats  = _latest_stats_per_client(stats_history)

        # Klanten met koppeling eerst, daarna de rest
        gekoppeld   = [c for c in clients if c.get("klant_id") in latest_stats]
        ongekoppeld = [c for c in clients if c.get("klant_id") not in latest_stats]
        sorted_clients = gekoppeld + ongekoppeld

        options = {c["bedrijfsnaam"]: c for c in sorted_clients}
        labels = [
            f"{naam} {'' if c.get('klant_id') in latest_stats else '— nog niet gekoppeld'}".strip()
            for naam, c in options.items()
        ]
        label_to_client = dict(zip(labels, options.values()))

        selected_label = st.selectbox("Selecteer een klant", labels)
        selected_client = label_to_client[selected_label]
        klant_id = selected_client.get("klant_id", "")

        if klant_id not in latest_stats:
            st.info(
                f"📭 **{selected_client['bedrijfsnaam']}** is nog niet gekoppeld aan het "
                "Meta Business Manager-portfolio. Statistieken verschijnen hier zodra de "
                "koppeling is gemaakt (zie blueprint *meta_insights_koppeling*)."
            )
        else:
            latest = latest_stats[klant_id]
            history = sorted(
                [r for r in stats_history if r.get("klant_id") == klant_id],
                key=lambda r: r.get("datum", ""),
            )

            with st.container(key=f"statcard-overview-{klant_id}"):
                st.markdown(
                    f'<p class="ts-stat-title">📅 Laatste meting: {latest.get("datum", "—")}</p>',
                    unsafe_allow_html=True,
                )

                heeft_ig = latest.get("instagram_volgers") is not None
                heeft_fb = latest.get("facebook_volgers") is not None

                if heeft_ig:
                    st.markdown(
                        f'<p style="font-size:13px;font-weight:700;color:{PLATFORM_COLORS["instagram"]};'
                        f'margin:14px 0 4px;">Instagram</p>',
                        unsafe_allow_html=True,
                    )
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Volgers", _format_stat(latest.get("instagram_volgers")),
                              _stat_delta(history, "instagram_volgers"))
                    c2.metric("Bereik (7d)", _format_stat(latest.get("instagram_bereik_7d")),
                              _stat_delta(history, "instagram_bereik_7d"))
                    c3.metric("Weergaven (7d)", _format_stat(latest.get("instagram_impressies_7d")),
                              _stat_delta(history, "instagram_impressies_7d"))
                    c4.metric("Profielbezoeken (7d)", _format_stat(latest.get("instagram_profielbezoeken_7d")),
                              _stat_delta(history, "instagram_profielbezoeken_7d"))

                if heeft_fb:
                    st.markdown(
                        f'<p style="font-size:13px;font-weight:700;color:{PLATFORM_COLORS["facebook"]};'
                        f'margin:18px 0 4px;">Facebook</p>',
                        unsafe_allow_html=True,
                    )
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Volgers", _format_stat(latest.get("facebook_volgers")),
                              _stat_delta(history, "facebook_volgers"))
                    c2.metric("Bereik (7d)", _format_stat(latest.get("facebook_bereik_7d")),
                              _stat_delta(history, "facebook_bereik_7d"))
                    c3.metric("Paginaweergaven (7d)", _format_stat(latest.get("facebook_impressies_7d")),
                              _stat_delta(history, "facebook_impressies_7d"))
                    c4.metric("Engagement (7d)", _format_stat(latest.get("facebook_engagement_7d")),
                              _stat_delta(history, "facebook_engagement_7d"))

                # ── Trendgrafiek (zodra er meerdere metingen zijn) ──────────────
                if len(history) > 1:
                    st.markdown(
                        '<p class="ts-stat-title" style="margin-top:18px;">Volgersgroei over tijd</p>',
                        unsafe_allow_html=True,
                    )
                    chart_cols = {"datum": [r["datum"] for r in history]}
                    if heeft_ig:
                        chart_cols["Instagram volgers"] = [r.get("instagram_volgers") for r in history]
                    if heeft_fb:
                        chart_cols["Facebook volgers"] = [r.get("facebook_volgers") for r in history]
                    if len(chart_cols) > 1:
                        import pandas as pd
                        df = pd.DataFrame(chart_cols).set_index("datum")
                        st.line_chart(df)
                else:
                    st.caption(
                        "📈 De trendgrafiek verschijnt zodra er meerdere metingen zijn opgeslagen "
                        "(de statistieken worden periodiek opnieuw opgehaald)."
                    )

            # ── Volgersgroei (laatste 7 dagen) ──────────────────────────────────
            with st.container(key=f"statcard-growth-{klant_id}"):
                st.markdown('<p class="ts-stat-title">📈 Volgersgroei</p>', unsafe_allow_html=True)
                ig_growth = _followers_growth(history, "instagram_volgers") if heeft_ig else None
                fb_growth = _followers_growth(history, "facebook_volgers") if heeft_fb else None

                if ig_growth is None and fb_growth is None:
                    st.markdown(
                        '<p class="ts-stat-caption">Nog niet genoeg metingen voor een groei-analyse.</p>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        '<p class="ts-stat-caption">Volgersverschil over de laatste 7 dagen '
                        '(of sinds de eerste meting)</p>',
                        unsafe_allow_html=True,
                    )
                    growth_items = []
                    if heeft_ig:
                        growth_items.append(("Instagram volgers", ig_growth))
                    if heeft_fb:
                        growth_items.append(("Facebook volgers", fb_growth))

                    growth_cols = st.columns(len(growth_items))
                    for col, (label, growth) in zip(growth_cols, growth_items):
                        if growth is None:
                            col.metric(label, "—")
                            continue
                        diff, pct = growth
                        sign = "+" if diff > 0 else ""
                        value_str = f"{sign}{int(diff)}"
                        delta_str = f"{sign}{pct:.1f}%" if pct is not None else None
                        col.metric(label, value_str, delta_str)

            # ── Posts, engagement, beste posttijden, demografie & AI-samenvatting ──
            posts_all = load_post_insights()
            all_posts = [p for p in posts_all if p.get("klant_id") == klant_id]

            if not all_posts:
                with st.container(key=f"statcard-noposts-{klant_id}"):
                    st.markdown('<p class="ts-stat-title">📭 Geen post-statistieken</p>', unsafe_allow_html=True)
                    st.caption(
                        "Nog geen post-statistieken opgehaald voor deze klant "
                        "(`systems/fetch_post_insights.py`)."
                    )
            else:
                today_d = date.today()
                cutoff_d = today_d - timedelta(days=31)
                period_label = f"{cutoff_d.strftime('%d-%m-%Y')} t/m {today_d.strftime('%d-%m-%Y')}"

                posts_31 = [
                    p for p in all_posts
                    if p.get("post_datum") and p["post_datum"] >= cutoff_d.isoformat()
                ]
                # Val terug op alle opgehaalde posts als er niets binnen 31 dagen valt
                posts = posts_31 if posts_31 else all_posts

                # ── Gemiddeld bereik (#2) ────────────────────────────────────────
                with st.container(key=f"statcard-engagement-{klant_id}"):
                    st.markdown('<p class="ts-stat-title">📈 Gemiddeld bereik per post</p>', unsafe_allow_html=True)
                    st.markdown(
                        '<p class="ts-stat-caption">Aantal mensen dat de post zag — '
                        f'laatste 31 dagen ({period_label})</p>',
                        unsafe_allow_html=True,
                    )
                    eng_cols = st.columns(2)
                    for i, platform in enumerate(["instagram", "facebook"]):
                        bereiken = [p["bereik"] for p in posts
                                    if p.get("platform") == platform and p.get("bereik") is not None]
                        if bereiken:
                            avg_bereik = sum(bereiken) / len(bereiken)
                            eng_cols[i].metric(
                                f"{platform.capitalize()} ({len(bereiken)} posts)",
                                _format_stat(avg_bereik),
                            )
                        else:
                            eng_cols[i].metric(f"{platform.capitalize()}", "—")

                # ── Top-presterende posts (#1) ──────────────────────────────────
                with st.container(key=f"statcard-topposts-{klant_id}"):
                    st.markdown('<p class="ts-stat-title">🏆 Best presterende posts</p>', unsafe_allow_html=True)
                    fallback_note = "" if posts_31 else " — geen posts in de laatste 31 dagen, toont recentste posts"
                    st.markdown(
                        f'<p class="ts-stat-caption">Periode: laatste 31 dagen ({period_label}){fallback_note}</p>',
                        unsafe_allow_html=True,
                    )

                    ranked = sorted(
                        [p for p in posts if p.get("engagement_rate") is not None],
                        key=lambda p: p["engagement_rate"], reverse=True,
                    )
                    if ranked:
                        platform_icons = {"instagram": "📷", "facebook": "📘"}
                        for i, p in enumerate(ranked[:5], start=1):
                            icon = platform_icons.get(p.get("platform", ""), "📱")
                            datum_fmt = p.get("post_datum", "")
                            try:
                                datum_fmt = datetime.strptime(p["post_datum"], "%Y-%m-%d").strftime("%d-%m-%Y")
                            except (ValueError, KeyError):
                                pass
                            bereik = p.get("bereik")
                            interacties = p.get("interacties")
                            link = p.get("link", "")
                            link_html = (
                                f'<a class="ts-post-link" href="{link}" target="_blank" title="Bekijk post">↗︎</a>'
                                if link else ""
                            )
                            afbeelding_url = p.get("afbeelding_url", "")
                            thumb_html = (
                                f'<img class="ts-post-thumb" src="{afbeelding_url}" alt="">'
                                if afbeelding_url else '<div class="ts-post-thumb ts-post-thumb-empty">{}</div>'.format(icon)
                            )
                            st.markdown(
                                f'<div class="ts-post-card">'
                                f'<div class="ts-post-rank">#{i}</div>'
                                f'{thumb_html}'
                                f'<div class="ts-post-body">'
                                f'<div class="ts-post-meta">{icon} {p.get("platform", "").capitalize()} · {datum_fmt}</div>'
                                f'<div class="ts-post-caption">{p.get("caption_kort", "(geen tekst)")}</div>'
                                f'</div>'
                                f'<div class="ts-post-stats">'
                                f'<div class="ts-post-stat"><div class="ts-post-stat-value">{_format_stat(bereik) if bereik is not None else "—"}</div>'
                                f'<div class="ts-post-stat-label">Bereik</div></div>'
                                f'<div class="ts-post-stat"><div class="ts-post-stat-value">{_format_stat(interacties) if interacties is not None else "—"}</div>'
                                f'<div class="ts-post-stat-label">Interacties</div></div>'
                                f'</div>'
                                f'{link_html}'
                                f'</div>',
                                unsafe_allow_html=True,
                            )
                    else:
                        st.caption("Nog geen posts met bereik-data beschikbaar.")

                # ── Beste momenten om te posten (#6) ────────────────────────────
                with st.container(key=f"statcard-bestdays-{klant_id}"):
                    st.markdown('<p class="ts-stat-title">🕐 Beste dagen om te posten</p>', unsafe_allow_html=True)
                    st.markdown(
                        '<p class="ts-stat-caption">Gemiddeld bereik (weergaven) per weekdag '
                        '(o.b.v. de meest recent opgehaalde posts)</p>',
                        unsafe_allow_html=True,
                    )

                    day_bereik: dict[str, list[float]] = {}
                    for p in all_posts:
                        if p.get("bereik") is None or not p.get("post_datum"):
                            continue
                        try:
                            weekday = datetime.strptime(p["post_datum"], "%Y-%m-%d").weekday()
                        except ValueError:
                            continue
                        day_bereik.setdefault(DAGEN_NL[weekday], []).append(p["bereik"])

                    if day_bereik:
                        import pandas as pd
                        avg_per_day = {dag: sum(v) / len(v) for dag, v in day_bereik.items()}
                        df_days = pd.DataFrame(
                            {"Gem. bereik": avg_per_day}
                        ).reindex(DAGEN_NL).dropna()
                        st.bar_chart(df_days, color=BRAND["primary"])
                        best_dag = max(avg_per_day, key=avg_per_day.get)
                        st.caption(f"📌 Op basis van de afgelopen posts presteert **{best_dag}** gemiddeld het best qua bereik.")
                    else:
                        st.caption("Nog niet genoeg data om beste posttijden te bepalen.")

                # ── Volgers-demografie (#7) ─────────────────────────────────────
                with st.container(key=f"statcard-demo-{klant_id}"):
                    st.markdown('<p class="ts-stat-title">👥 Doelgroep (Instagram)</p>', unsafe_allow_html=True)
                    demo = load_demographics()
                    demo_per_dim = _latest_demo_per_client(demo, klant_id)
                    if not demo_per_dim:
                        st.caption(
                            "Nog geen demografische gegevens beschikbaar (vereist o.a. minimaal "
                            "~100 Instagram-volgers)."
                        )
                    else:
                        import pandas as pd
                        GESLACHT_LABELS = {"F": "Vrouw", "M": "Man", "U": "Onbekend"}
                        demo_cols = st.columns(len(demo_per_dim))
                        for i, (dim, rows) in enumerate(demo_per_dim.items()):
                            title = "Leeftijd / geslacht" if dim == "leeftijd_geslacht" else "Man / vrouw"
                            labels = {
                                GESLACHT_LABELS.get(r["waarde"], r["waarde"]) if dim == "geslacht" else r["waarde"]: r["aantal"]
                                for r in rows
                            }
                            df_demo = pd.DataFrame(
                                labels.items(),
                                columns=["categorie", "aantal"],
                            ).set_index("categorie").sort_values("aantal", ascending=False).head(10)
                            with demo_cols[i]:
                                st.caption(title)
                                st.bar_chart(df_demo, color=BRAND["primary"])

                # ── AI-weeksamenvatting (#4) ────────────────────────────────────
                with st.container(key=f"statcard-aisummary-{klant_id}"):
                    st.markdown('<p class="ts-stat-title">🤖 AI-samenvatting</p>', unsafe_allow_html=True)
                    previous = history[-2] if len(history) > 1 else {}
                    summary = generate_ai_summary(
                        klant_id, selected_client["bedrijfsnaam"], latest, previous, ranked,
                    )
                    if summary:
                        st.info(summary)
                    else:
                        st.caption("AI-samenvatting niet beschikbaar (controleer ANTHROPIC_API_KEY).")

    st.caption(f"Laatste update: {_now_ams().strftime('%H:%M:%S')}")


# ── Team-beheer (alleen admin) ────────────────────────────────────────────────
if tab_team is not None:
    with tab_team:
        st.subheader("👥 Medewerkers & klant-toewijzingen")
        st.caption("Wijs per medewerker de klanten toe die zij mogen zien en beoordelen. "
                   "Wijzigingen gaan direct in op het volgende inlogmoment van de medewerker.")

        _t_sid = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID") or st.secrets.get("GOOGLE_SHEETS_SPREADSHEET_ID", "")
        _t_saj = _sa_info_json()

        if not _t_sid or not _t_saj:
            st.error("Credentials ontbreken.")
        else:
            # Lijst van alle medewerkers uit secrets (excl. admins)
            _all_users = {
                uname: udata.get("name", uname)
                for uname, udata in st.secrets.get("auth", {}).get("users", {}).items()
            }

            if not _all_users:
                st.info("Geen medewerkers gevonden in de secrets. "
                        "Voeg medewerkers toe via Streamlit Cloud → Secrets.")
            else:
                _current_assignments = load_medewerker_assignments(_t_sid, _t_saj)
                _all_client_names    = sorted(c.get("bedrijfsnaam", "") for c in _all_clients)

                st.divider()
                _changed = False
                _new_assignments = dict(_current_assignments)

                for _uname, _display_name in sorted(_all_users.items(), key=lambda x: x[1]):
                    current_sel = _current_assignments.get(_uname, [])
                    if current_sel == "ALL":
                        current_sel = _all_client_names

                    col_name, col_select = st.columns([1, 3])
                    with col_name:
                        st.markdown(f"**{_display_name}**")
                        st.caption(f"@{_uname}")
                    with col_select:
                        new_sel = st.multiselect(
                            "Klanten",
                            options=_all_client_names,
                            default=[c for c in current_sel if c in _all_client_names],
                            key=f"team_sel_{_uname}",
                            label_visibility="collapsed",
                            placeholder="Selecteer klanten...",
                        )
                    _new_assignments[_uname] = new_sel
                    if set(new_sel) != set(current_sel if isinstance(current_sel, list) else _all_client_names):
                        _changed = True

                    st.markdown(
                        "<hr style='margin:6px 0;border:none;border-top:1px solid #f0f0f0;'>",
                        unsafe_allow_html=True,
                    )

                if st.button("💾 Toewijzingen opslaan", type="primary", use_container_width=False):
                    with st.spinner("Opslaan..."):
                        save_medewerker_assignments(_t_sid, _t_saj, _new_assignments)
                    st.success("✅ Toewijzingen opgeslagen!")
